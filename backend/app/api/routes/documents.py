"""
Document Upload API Routes — P3: RAG Project Isolation.

Allows users to upload PDF/DOCX/TXT files to a project's RAG context.
Documents are chunked and ingested into ChromaDB with project_id metadata
so that agents only see context from their own project.
"""
import os
import logging
from pathlib import Path
from typing import List

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, status
from sqlalchemy.orm import Session

from app.database import get_db
from app.config import settings
from app.models.project import Project
from app.models.project_document import ProjectDocument
from app.models.user import User
from app.api.routes.auth import get_current_user

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/projects", tags=["Documents"])

# Allowed file extensions and their MIME types
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".csv"}
MAX_FILE_SIZE = 20 * 1024 * 1024  # 20 MB

# Upload directory: <UPLOAD_DIR>/project_documents/<project_id>/
DOCUMENTS_DIR = Path(settings.UPLOAD_DIR) / "project_documents"


def _get_project_or_404(project_id: int, user: User, db: Session) -> Project:
    """Get project belonging to current user or raise 404."""
    project = db.query(Project).filter(
        Project.id == project_id,
        Project.user_id == user.id,
    ).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    return project


def _chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    """Split text into overlapping chunks."""
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start = end - overlap
    return [c for c in chunks if c.strip()]


def _extract_text(file_path: Path, content_type: str) -> str:
    """Extract plain text from uploaded file."""
    suffix = file_path.suffix.lower()

    if suffix == ".txt" or suffix == ".md" or suffix == ".csv":
        return file_path.read_text(encoding="utf-8", errors="replace")

    if suffix == ".pdf":
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(str(file_path))
            text_parts = []
            for page in doc:
                text_parts.append(page.get_text())
            doc.close()
            return "\n".join(text_parts)
        except ImportError:
            logger.warning("PyMuPDF (fitz) not installed, falling back to basic PDF read")
            return file_path.read_text(encoding="utf-8", errors="replace")

    if suffix == ".docx":
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(str(file_path))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except ImportError:
            logger.warning("python-docx not installed, falling back to basic read")
            return file_path.read_text(encoding="utf-8", errors="replace")

    return file_path.read_text(encoding="utf-8", errors="replace")


# ========================================
# ENDPOINTS
# ========================================

@router.post("/{project_id}/documents")
def upload_document(
    project_id: int,
    file: UploadFile = File(...),
    collection_name: str = "technical",
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Upload a document to a project's RAG context.

    The file is saved, chunked, and ingested into ChromaDB with
    project_id metadata for isolation.
    """
    project = _get_project_or_404(project_id, current_user, db)

    # Validate file extension
    ext = Path(file.filename or "").suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"File type '{ext}' not allowed. Allowed: {', '.join(ALLOWED_EXTENSIONS)}",
        )

    # Validate collection name
    from app.services.rag_service import COLLECTIONS
    if collection_name not in COLLECTIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Unknown collection '{collection_name}'. Available: {', '.join(COLLECTIONS.keys())}",
        )

    # Read file content
    content = file.file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail=f"File too large (max {MAX_FILE_SIZE // 1024 // 1024} MB)")

    # Save to disk
    project_dir = DOCUMENTS_DIR / str(project_id)
    project_dir.mkdir(parents=True, exist_ok=True)
    file_path = project_dir / file.filename
    file_path.write_bytes(content)

    # Create DB record
    doc = ProjectDocument(
        project_id=project_id,
        filename=file.filename,
        file_path=str(file_path),
        file_size=len(content),
        content_type=file.content_type,
        collection_name=collection_name,
        status="processing",
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)

    # Extract text and chunk
    try:
        text = _extract_text(file_path, file.content_type or "")
        if not text.strip():
            doc.status = "error"
            doc.error_message = "No text could be extracted from file"
            db.commit()
            return {
                "id": doc.id,
                "filename": doc.filename,
                "status": "error",
                "error": doc.error_message,
            }

        chunks = _chunk_text(text)

        # Ingest into ChromaDB
        from app.services.rag_service import ingest_document
        chunk_count = ingest_document(
            collection_name=collection_name,
            chunks=chunks,
            metadata={"source": file.filename, "project_id": str(project_id)},
            project_id=project_id,
            document_id=doc.id,
        )

        doc.chunk_count = chunk_count
        doc.status = "ready"
        db.commit()

        logger.info(f"Document '{file.filename}' ingested: {chunk_count} chunks into {collection_name} for project {project_id}")

    except Exception as e:
        logger.error(f"Error ingesting document '{file.filename}': {e}")
        doc.status = "error"
        doc.error_message = str(e)[:500]
        db.commit()
        return {
            "id": doc.id,
            "filename": doc.filename,
            "status": "error",
            "error": str(e)[:200],
        }

    return {
        "id": doc.id,
        "filename": doc.filename,
        "file_size": doc.file_size,
        "collection_name": collection_name,
        "chunk_count": chunk_count,
        "status": "ready",
    }


@router.get("/{project_id}/documents")
def list_documents(
    project_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """List all documents uploaded to a project."""
    _get_project_or_404(project_id, current_user, db)

    docs = (
        db.query(ProjectDocument)
        .filter(ProjectDocument.project_id == project_id)
        .order_by(ProjectDocument.created_at.desc())
        .all()
    )

    return [
        {
            "id": d.id,
            "filename": d.filename,
            "file_size": d.file_size,
            "collection_name": d.collection_name,
            "chunk_count": d.chunk_count,
            "status": d.status,
            "created_at": d.created_at.isoformat() if d.created_at else None,
        }
        for d in docs
    ]


@router.delete("/{project_id}/documents/{document_id}")
def delete_document(
    project_id: int,
    document_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Delete a document and its chunks from ChromaDB."""
    _get_project_or_404(project_id, current_user, db)

    doc = db.query(ProjectDocument).filter(
        ProjectDocument.id == document_id,
        ProjectDocument.project_id == project_id,
    ).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    # Remove chunks from ChromaDB
    deleted_chunks = 0
    if doc.status == "ready" and doc.chunk_count and doc.chunk_count > 0:
        try:
            from app.services.rag_service import delete_project_document_chunks
            deleted_chunks = delete_project_document_chunks(doc.collection_name, document_id)
        except Exception as e:
            logger.error(f"Error deleting chunks for document {document_id}: {e}")

    # Remove file from disk
    try:
        file_path = Path(doc.file_path)
        if file_path.exists():
            file_path.unlink()
    except Exception as e:
        logger.warning(f"Could not delete file {doc.file_path}: {e}")

    # Remove DB record
    db.delete(doc)
    db.commit()

    logger.info(f"Document '{doc.filename}' deleted: {deleted_chunks} chunks removed from {doc.collection_name}")

    return {
        "deleted": True,
        "filename": doc.filename,
        "chunks_removed": deleted_chunks,
    }
