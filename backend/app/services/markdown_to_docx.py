"""
Markdown to DOCX Converter - Enhanced with Mermaid Image Support
Converts Emma's SDS markdown to professional DOCX format.
"""

import os
import re
import subprocess
import tempfile
from datetime import datetime
from typing import Optional, List

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

PUPPETEER_CONFIG = "/tmp/puppeteer-config.json"

def _ensure_puppeteer_config():
    if not os.path.exists(PUPPETEER_CONFIG):
        with open(PUPPETEER_CONFIG, 'w') as f:
            f.write('{"args": ["--no-sandbox", "--disable-setuid-sandbox"]}')

COLORS = {"primary": "003366", "secondary": "006699", "accent": "0099CC", "alt_row": "F5F5F5"}
FONT_COLORS = {
    "primary": RGBColor(0, 51, 102), "secondary": RGBColor(0, 102, 153),
    "accent": RGBColor(0, 153, 204), "text": RGBColor(51, 51, 51),
    "muted": RGBColor(128, 128, 128), "white": RGBColor(255, 255, 255),
}


def convert_mermaid_to_image(mermaid_code: str, output_dir: str = "/tmp") -> Optional[str]:
    """Convert Mermaid diagram to PNG using mmdc."""
    _ensure_puppeteer_config()
    try:
        fixed_code = _fix_mermaid_syntax(mermaid_code)
        with tempfile.NamedTemporaryFile(mode='w', suffix='.mmd', delete=False) as f:
            f.write(fixed_code)
            mmd_path = f.name
        
        png_path = os.path.join(output_dir, f"diagram_{os.path.basename(mmd_path).replace('.mmd', '')}.png")
        result = subprocess.run(
            ['mmdc', '-i', mmd_path, '-o', png_path, '-b', 'white', '-w', '800', '-p', PUPPETEER_CONFIG],
            capture_output=True, timeout=60
        )
        os.unlink(mmd_path)
        
        if result.returncode == 0 and os.path.exists(png_path):
            return png_path
        print(f"⚠️ Mermaid conversion failed: {(result.stderr or result.stdout).decode()[:200]}")
        return None
    except Exception as e:
        print(f"⚠️ Mermaid error: {e}")
        return None


def _fix_mermaid_syntax(code: str) -> str:
    """Fix erDiagram syntax for newer mermaid versions."""
    lines = code.split('\n')
    fixed_lines = []
    in_entity = False
    
    for line in lines:
        stripped = line.strip()
        if re.match(r'^\w+\s*\{', stripped):
            in_entity = True
            fixed_lines.append(line)
            continue
        if stripped == '}':
            in_entity = False
            fixed_lines.append(line)
            continue
        
        if in_entity and stripped and not stripped.startswith('%%'):
            # Convert "FieldName Type" to "type FieldName"
            field_match = re.match(r'^(\w+)\s+(\w+)(?:\s+(.*))?$', stripped)
            if field_match:
                field_name, field_type, extra = field_match.groups()
                type_map = {'PK': 'string', 'FK': 'string', 'string': 'string', 'text': 'string',
                           'number': 'decimal', 'decimal': 'decimal', 'boolean': 'boolean', 'date': 'date'}
                if field_type.upper() in ['PK', 'FK'] or field_type.lower() in type_map:
                    mapped_type = type_map.get(field_type.upper(), type_map.get(field_type.lower(), 'string'))
                    fixed_lines.append(f"        {mapped_type} {field_name}" + (f" {extra}" if extra else ""))
                    continue
        fixed_lines.append(line)
    return '\n'.join(fixed_lines)


def set_cell_shading(cell, hex_color: str):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(shading)


def create_professional_table(doc: Document, headers: List[str], rows: List[List[str]], title: str = None):
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = FONT_COLORS["primary"]
    
    num_cols = len(headers) if headers else (len(rows[0]) if rows else 1)
    num_rows = (1 if headers else 0) + len(rows)
    if num_rows == 0 or num_cols == 0:
        return
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    if headers:
        for i, header in enumerate(headers):
            if i < len(table.rows[0].cells):
                cell = table.rows[0].cells[i]
                cell.text = str(header)
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.bold = True
                        run.font.size = Pt(10)
                        run.font.color.rgb = FONT_COLORS["white"]
                set_cell_shading(cell, COLORS["primary"])
    
    start_row = 1 if headers else 0
    for row_idx, row_data in enumerate(rows):
        if start_row + row_idx >= len(table.rows):
            break
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < len(table.rows[start_row + row_idx].cells):
                cell = table.rows[start_row + row_idx].cells[col_idx]
                cell.text = str(cell_text) if cell_text else ""
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)
                if row_idx % 2 == 1:
                    set_cell_shading(cell, COLORS["alt_row"])
    doc.add_paragraph()


def convert_markdown_to_docx(markdown_content: str, output_path: str, project_name: str = "Project",
                              client_name: str = "", version: str = "1.0") -> str:
    if not DOCX_AVAILABLE:
        md_path = output_path.replace('.docx', '.md')
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        return md_path
    
    doc = Document()
    _setup_document(doc)
    _add_title_page(doc, project_name, client_name, version)
    _parse_markdown_enhanced(doc, markdown_content)
    doc.save(output_path)
    return output_path


def _setup_document(doc: Document):
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)


def _add_title_page(doc: Document, project_name: str, client_name: str, version: str):
    for _ in range(3):
        doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SOLUTION DESIGN SPECIFICATION")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = FONT_COLORS["primary"]
    
    doc.add_paragraph()
    project = doc.add_paragraph()
    project.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = project.add_run(project_name)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = FONT_COLORS["secondary"]
    
    if client_name:
        doc.add_paragraph()
        client = doc.add_paragraph()
        client.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = client.add_run(f"Client: {client_name}")
        run.font.size = Pt(14)
    
    for _ in range(5):
        doc.add_paragraph()
    
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Version {version}\n").font.size = Pt(12)
    meta.add_run(f"{datetime.now().strftime('%d %B %Y')}\n").font.size = Pt(12)
    meta.add_run("\n")
    run = meta.add_run("Document généré par Digital Humans Platform")
    run.font.size = Pt(10)
    run.font.color.rgb = FONT_COLORS["muted"]
    
    doc.add_page_break()


def _parse_markdown_enhanced(doc: Document, markdown: str):
    lines = markdown.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i]
        
        # Mermaid code blocks - convert to image
        if line.strip().startswith('```mermaid'):
            mermaid_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                mermaid_lines.append(lines[i])
                i += 1
            
            mermaid_code = '\n'.join(mermaid_lines)
            if mermaid_code.strip():
                img_path = convert_mermaid_to_image(mermaid_code)
                if img_path and os.path.exists(img_path):
                    # Add diagram as image
                    doc.add_paragraph()
                    para = doc.add_paragraph()
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = para.add_run()
                    run.add_picture(img_path, width=Inches(5.5))
                    doc.add_paragraph()
                    try:
                        os.unlink(img_path)
                    except:
                        pass
                else:
                    # Fallback: format as entity tables
                    _format_mermaid_as_tables(doc, mermaid_code)
            i += 1
            continue
        
        # Other code blocks
        if line.strip().startswith('```'):
            lang = line.strip()[3:]
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                if lang:
                    p = doc.add_paragraph()
                    run = p.add_run(f"Code ({lang}):")
                    run.bold = True
                    run.font.size = Pt(9)
                code_para = doc.add_paragraph()
                run = code_para.add_run('\n'.join(code_lines))
                run.font.name = 'Courier New'
                run.font.size = Pt(8)
            i += 1
            continue
        
        # Tables
        if line.strip().startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            if table_lines:
                _parse_and_create_table(doc, table_lines)
            continue
        
        # Headings
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=4)
        elif line.strip() in ['---', '***', '___']:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("─" * 50)
            run.font.color.rgb = FONT_COLORS["muted"]
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            para = doc.add_paragraph(style='List Bullet')
            _add_formatted_text(para, line.strip()[2:])
        elif re.match(r'^\d+\.\s', line.strip()):
            para = doc.add_paragraph(style='List Number')
            _add_formatted_text(para, re.sub(r'^\d+\.\s', '', line.strip()))
        elif line.strip():
            para = doc.add_paragraph()
            _add_formatted_text(para, line)
        
        i += 1


def _format_mermaid_as_tables(doc: Document, mermaid_code: str):
    """Fallback: parse erDiagram and create tables."""
    label = doc.add_paragraph()
    run = label.add_run("📊 Diagramme ERD:")
    run.bold = True
    run.font.color.rgb = FONT_COLORS["secondary"]
    
    entities = {}
    relationships = []
    current_entity = None
    
    for line in mermaid_code.split('\n'):
        line = line.strip()
        entity_match = re.match(r'^(\w+)\s*\{', line)
        if entity_match:
            current_entity = entity_match.group(1)
            entities[current_entity] = []
            continue
        if line == '}':
            current_entity = None
            continue
        if current_entity and line:
            field_match = re.match(r'^(\w+)\s+(\w+)', line)
            if field_match:
                entities[current_entity].append((field_match.group(1), field_match.group(2)))
        rel_match = re.match(r'^(\w+)\s*\|.*\|\s*(\w+)\s*:\s*(.+)$', line)
        if rel_match:
            relationships.append((rel_match.group(1), rel_match.group(2), rel_match.group(3).strip('"')))
    
    for name, fields in entities.items():
        if fields:
            create_professional_table(doc, ["Champ", "Type"], [[f[0], f[1]] for f in fields], f"📦 {name}")
    
    if relationships:
        create_professional_table(doc, ["Parent", "Enfant", "Relation"], relationships, "🔗 Relations")


def _parse_and_create_table(doc: Document, table_lines: List[str]):
    headers = []
    rows = []
    for i, line in enumerate(table_lines):
        if re.match(r'^\|[\s\-:\|]+\|$', line.strip()):
            continue
        cells = [c.strip() for c in line.strip().split('|')[1:-1]]
        if not headers:
            headers = cells
        else:
            rows.append(cells)
    if headers or rows:
        create_professional_table(doc, headers, rows)


def _add_formatted_text(para, text: str):
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)'
    parts = re.split(pattern, text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = para.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = para.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = FONT_COLORS["secondary"]
        else:
            para.add_run(part)


if __name__ == "__main__":
    test_md = """# Solution Design Specification
## Customer Feedback System

### 1. Executive Summary
This document presents the **solution design** for the *Customer Feedback Management* system.

### 2. Data Model

```mermaid
erDiagram
    Account {
        Id PK
        Name string
        Average_Rating__c number
    }
    Customer_Feedback__c {
        Id PK
        Account__c FK
        Rating__c number
        Comments__c text
    }
    Account ||--o{ Customer_Feedback__c : has
```

### 3. Business Requirements

| ID | Titre | Priorité | Catégorie |
|-----|-------|----------|-----------|
| BR-001 | Capture feedback | MUST | Data Model |
| BR-002 | Calculate ratings | SHOULD | Automation |

### 4. Implementation
- Phase 1: Data model setup
- Phase 2: Automation development

---
**Document generated by Digital Humans Platform**
"""
    
    if DOCX_AVAILABLE:
        output = convert_markdown_to_docx(test_md, "/tmp/test_sds_with_diagram.docx", "Customer Feedback System", "Acme Corp")
        print(f"✅ Document created: {output} ({os.path.getsize(output)} bytes)")
    else:
        print("⚠️ python-docx not available")
