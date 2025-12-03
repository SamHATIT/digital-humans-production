"""
PM Orchestrator Service - Core orchestration logic
Executes agents sequentially and generates SDS document
"""
from app.services.document_generator import generate_professional_sds, ProfessionalDocumentGenerator
from app.services.sds_template_generator import generate_sds_from_template
import asyncio
import os
from typing import List, Dict, Any, Optional
from datetime import datetime, timezone
from sqlalchemy.orm import Session
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging

from app.utils.cost_calculator import calculate_cost

from app.models.project import Project
from app.models.execution import Execution, ExecutionStatus
from app.database import SessionLocal
from app.services.agent_integration import AgentIntegrationService, AgentExecutionError
from app.services.quality_gates import QualityGatesService, QualityGateStatus
from app.services.artifact_service import ArtifactService
from app.schemas.artifact import ArtifactCreate, ArtifactTypeEnum
import re

logger = logging.getLogger(__name__)

# V2: Maximum iterations for BA <-> Architect Q&A cycle
MAX_ITERATIONS = 3


class PMOrchestratorService:
    """Main orchestration service for PM workflow"""

    def __init__(self, db: Session):
        self.db = db
        self.agent_service = AgentIntegrationService()
        self.quality_service = QualityGatesService()
        self.artifact_service = ArtifactService(db)

    async def execute_agents(
        self,
        execution_id: int,
        project_id: int,
        selected_agents: List[str]
    ) -> Dict[str, Any]:
        """
        Main execution method - orchestrates all selected agents

        Args:
            execution_id: ID of the execution record
            project_id: ID of the project
            selected_agents: List of agent IDs to execute

        Returns:
            Dict with execution results
        """
        try:
            # Get project
            project = self.db.query(Project).filter(Project.id == project_id).first()
            if not project:
                raise ValueError(f"Project {project_id} not found")

            # Get execution
            execution = self.db.query(Execution).filter(Execution.id == execution_id).first()
            if not execution:
                raise ValueError(f"Execution {execution_id} not found")

            # Initialize status tracking
            agent_execution_status = {
                agent_id: {
                    "state": "waiting",
                    "progress": 0,
                    "message": "Waiting to start...",
                    "started_at": None,
                    "completed_at": None,
                    "error": None
                }
                for agent_id in selected_agents
            }

            # Update execution
            execution.status = ExecutionStatus.RUNNING

            # V2: Initialize validation gates
            try:
                self._initialize_validation_gates(execution_id)
            except Exception as e:
                logger.warning(f"Could not initialize validation gates: {e}")
            execution.started_at = datetime.now(timezone.utc)
            execution.agent_execution_status = agent_execution_status
            self.db.commit()

            # Agent order mapping
            AGENT_ORDER = {
                'ba': 1, 'architect': 2, 'apex': 3, 'lwc': 4,
                'admin': 5, 'qa': 6, 'devops': 7, 'data': 8,
                'trainer': 9, 'pm': 10
            }

            # Sort agents by order
            sorted_agents = sorted(selected_agents, key=lambda x: AGENT_ORDER.get(x, 99))
            # Remove PM from agent execution list (PM only generates final SDS)
            sorted_agents = [a for a in sorted_agents if a != "pm"]


            # Execute each agent
            agent_outputs = {}
            
            # V2: Check if we should use iteration mode (BA + Architect in sorted_agents)
            use_iteration_mode = "ba" in sorted_agents and "architect" in sorted_agents
            ba_architect_done = False

            for agent_id in sorted_agents:
                # V2: Handle BA and Architect together with iterations
                if use_iteration_mode and agent_id in ["ba", "architect"] and not ba_architect_done:
                    if agent_id == "ba":
                        # Execute BA + Architect with iteration
                        iteration_result = await self._handle_ba_architect_iterations(
                            project=project,
                            execution=execution,
                            agent_outputs=agent_outputs,
                            agent_execution_status=agent_execution_status
                        )
                        agent_outputs["ba"] = iteration_result["ba"]
                        agent_outputs["architect"] = iteration_result["architect"]
                        ba_architect_done = True
                        logger.info(f"[V2] BA-Architect iteration completed in {iteration_result['iterations']} iterations")
                        continue
                    elif agent_id == "architect":
                        # Already done in iteration with BA
                        continue
                try:
                    # Update status to running
                    agent_execution_status[agent_id]["state"] = "running"
                    agent_execution_status[agent_id]["started_at"] = datetime.now(timezone.utc).isoformat()
                    agent_execution_status[agent_id]["message"] = "Processing project requirements..."
                    agent_execution_status[agent_id]["progress"] = 10
                    execution.agent_execution_status = agent_execution_status
                    execution.current_agent = agent_id
                    self.db.commit()

                    # Execute agent
                    result = await self._execute_single_agent(
                        agent_id=agent_id,
                        project=project,
                        previous_outputs={aid: agent_outputs[aid].get("output", {}) for aid in agent_outputs},
                        execution=execution,
                        agent_status=agent_execution_status[agent_id]
                    )

                    # Store full result (includes output, quality_check, execution_time, etc.)
                    agent_outputs[agent_id] = result

                    # V2: Extract and persist artifacts from agent output
                    try:
                        extracted = self._extract_artifacts_from_agent_output(
                            agent_id=agent_id,
                            output=result,
                            execution_id=execution_id
                        )
                        if extracted:
                            self._persist_artifacts(extracted, execution_id)
                            # Check if gate should be updated
                            gate_info = self._check_gate_progress(execution_id, agent_id)
                            if gate_info:
                                self._update_gate_artifacts_count(execution_id, gate_info["gate_number"])
                    except Exception as e:
                        logger.warning(f"Artifact extraction failed for {agent_id}: {e}")

                    # Check if agent failed
                    if result.get("status") == "failed":
                        raise Exception(result.get("error", "Agent execution failed"))

                    # Update status to completed
                    agent_execution_status[agent_id]["state"] = "completed"
                    agent_execution_status[agent_id]["progress"] = 100
                    agent_execution_status[agent_id]["completed_at"] = datetime.now(timezone.utc).isoformat()
                    agent_execution_status[agent_id]["message"] = "Completed successfully"

                    # Log quality check summary
                    if "quality_check" in result:
                        qc = result["quality_check"]
                        logger.info(
                            f"Agent {agent_id} quality: {qc['overall_status']} "
                            f"({qc['passed']}/{qc['total_gates']} gates passed)"
                        )

                except Exception as e:
                    # Handle agent error
                    logger.error(f"Error executing agent {agent_id}: {str(e)}")
                    agent_execution_status[agent_id]["state"] = "error"
                    agent_execution_status[agent_id]["error"] = str(e)
                    agent_execution_status[agent_id]["message"] = f"Error: {str(e)[:100]}"
                    execution.agent_execution_status = agent_execution_status
                    self.db.commit()
                    raise

                finally:
                    execution.agent_execution_status = agent_execution_status
                    self.db.commit()


            # Aggregate token usage from all agents
            total_tokens = 0
            tokens_by_agent = {}
            for agent_id, output in agent_outputs.items():
                # Extract tokens from agent output metadata
                agent_output_data = output.get("output", {}).get("json_data", {})
                metadata = agent_output_data.get("metadata", {}) if agent_output_data else {}
                tokens_used = metadata.get("tokens_used", 0)
                model_used = metadata.get("model", "gpt-4o-mini")
                
                if tokens_used:
                    total_tokens += tokens_used
                    tokens_by_agent[agent_id] = {
                        "tokens": tokens_used,
                        "model": model_used,
                        "cost": calculate_cost(tokens_used, model_used)
                    }
                    logger.info(f"Agent {agent_id} used {tokens_used} tokens ({model_used})")
            
            # Calculate total cost
            total_cost = calculate_cost(total_tokens, "gpt-4o-mini")
            
            # Update execution with token tracking
            execution.total_tokens_used = total_tokens
            execution.total_cost = total_cost
            logger.info(f"Total execution tokens: {total_tokens}, estimated cost: ${total_cost:.4f}")
            self.db.commit()

            # Generate SDS document
            sds_path = await self._generate_sds_document(
                project=project,
                agent_outputs=agent_outputs,
                execution_id=execution_id
            )

            # Update execution as completed
            execution.status = ExecutionStatus.COMPLETED
            execution.completed_at = datetime.now(timezone.utc)
            execution.sds_document_path = sds_path
            
            # IMPORTANT: Update project status to completed
            project.status = "completed"
            logger.info(f"Project {project.id} marked as completed")
            execution.duration_seconds = int(
                (execution.completed_at - execution.started_at).total_seconds()
            )
            execution.current_agent = None
            self.db.commit()

            return {
                "success": True,
                "execution_id": execution_id,
                "sds_path": sds_path,
                "execution_time": execution.duration_seconds
            }

        except Exception as e:
            # Mark execution as failed
            if execution:
                execution.status = ExecutionStatus.FAILED
                execution.completed_at = datetime.now(timezone.utc)
                execution.current_agent = None
                self.db.commit()
            raise

    # ============ V2: BA <-> ARCHITECT ITERATION LOGIC ============
    
    async def _handle_ba_architect_iterations(
        self,
        project: "Project",
        execution: "Execution",
        agent_outputs: Dict[str, Any],
        agent_execution_status: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        Handle iterative refinement between BA and Architect.
        
        Flow:
        1. BA produces BR/UC
        2. Architect reviews and may ask questions (Q-xxx)
        3. BA answers questions
        4. Repeat until no questions or MAX_ITERATIONS
        5. Architect produces final ADR/SPEC
        
        Returns:
            Dict with ba_output and architect_output
        """
        iteration = 0
        questions_pending = True
        
        # First, execute BA
        logger.info(f"[V2 Iteration] Starting BA execution")
        agent_execution_status["ba"]["state"] = "running"
        agent_execution_status["ba"]["message"] = "Analyzing business requirements..."
        agent_execution_status["ba"]["started_at"] = datetime.now(timezone.utc).isoformat()
        execution.agent_execution_status = agent_execution_status
        execution.current_agent = "ba"
        self.db.commit()
        
        ba_result = await self._execute_single_agent(
            agent_id="ba",
            project=project,
            previous_outputs={},
            execution=execution,
            agent_status=agent_execution_status["ba"]
        )
        agent_outputs["ba"] = ba_result
        
        # Extract BA artifacts
        try:
            extracted = self._extract_artifacts_from_agent_output("ba", ba_result, execution.id)
            if extracted:
                self._persist_artifacts(extracted, execution.id)
        except Exception as e:
            logger.warning(f"BA artifact extraction failed: {e}")
        
        agent_execution_status["ba"]["state"] = "completed"
        agent_execution_status["ba"]["progress"] = 100
        agent_execution_status["ba"]["completed_at"] = datetime.now(timezone.utc).isoformat()
        agent_execution_status["ba"]["message"] = "Business analysis completed"
        
        # Now iterate with Architect
        while questions_pending and iteration < MAX_ITERATIONS:
            iteration += 1
            logger.info(f"[V2 Iteration] Architect iteration {iteration}/{MAX_ITERATIONS}")
            
            # Update status
            iteration_msg = f"Iteration {iteration}/{MAX_ITERATIONS}" if iteration > 1 else "Designing architecture..."
            agent_execution_status["architect"]["state"] = "running"
            agent_execution_status["architect"]["message"] = iteration_msg
            agent_execution_status["architect"]["started_at"] = datetime.now(timezone.utc).isoformat()
            agent_execution_status["architect"]["progress"] = 10 + (iteration - 1) * 20
            execution.agent_execution_status = agent_execution_status
            execution.current_agent = "architect"
            self.db.commit()
            
            # Build context with BA artifacts and any previous Q&A
            previous_outputs = {"ba": agent_outputs["ba"].get("output", {})}
            if "architect" in agent_outputs:
                previous_outputs["architect"] = agent_outputs["architect"].get("output", {})
            
            # Add answered questions context if any
            answered_questions = self._get_answered_questions(execution.id)
            if answered_questions:
                previous_outputs["answered_questions"] = answered_questions
            
            # Execute Architect
            architect_result = await self._execute_single_agent(
                agent_id="architect",
                project=project,
                previous_outputs=previous_outputs,
                execution=execution,
                agent_status=agent_execution_status["architect"]
            )
            agent_outputs["architect"] = architect_result
            
            # Extract Architect artifacts
            try:
                extracted = self._extract_artifacts_from_agent_output("architect", architect_result, execution.id)
                if extracted:
                    self._persist_artifacts(extracted, execution.id)
                    
                    # Check for questions in artifacts
                    questions = [a for a in extracted if a.get("artifact_type") == "question" or 
                                 a.get("artifact_code", "").startswith("Q-")]
                    
                    if questions and iteration < MAX_ITERATIONS:
                        # Architect has questions - BA must answer
                        logger.info(f"[V2 Iteration] Architect asked {len(questions)} questions")
                        
                        # Store questions
                        self._store_questions(questions, execution.id)
                        
                        # Have BA answer
                        ba_answer_result = await self._ba_answer_questions(
                            project=project,
                            questions=questions,
                            ba_artifacts=agent_outputs["ba"],
                            execution=execution,
                            agent_status=agent_execution_status["ba"]
                        )
                        
                        # Update BA output with answers
                        if ba_answer_result.get("success"):
                            agent_outputs["ba_answers"] = ba_answer_result
                            self._update_questions_with_answers(ba_answer_result, execution.id)
                    else:
                        questions_pending = False
                else:
                    questions_pending = False
                    
            except Exception as e:
                logger.warning(f"Architect artifact extraction failed: {e}")
                questions_pending = False
        
        # Final status update
        agent_execution_status["architect"]["state"] = "completed"
        agent_execution_status["architect"]["progress"] = 100
        agent_execution_status["architect"]["completed_at"] = datetime.now(timezone.utc).isoformat()
        agent_execution_status["architect"]["message"] = f"Architecture completed (iterations: {iteration})"
        execution.agent_execution_status = agent_execution_status
        self.db.commit()
        
        logger.info(f"[V2 Iteration] Completed with {iteration} iteration(s)")
        
        return {
            "ba": agent_outputs.get("ba"),
            "architect": agent_outputs.get("architect"),
            "iterations": iteration
        }
    
    def _get_answered_questions(self, execution_id: int) -> List[Dict]:
        """Get questions that have been answered"""
        try:
            from app.models.artifact import ExecutionArtifact
            questions = self.db.query(ExecutionArtifact).filter(
                ExecutionArtifact.execution_id == execution_id,
                ExecutionArtifact.artifact_type == "question"
            ).all()
            
            answered = []
            for q in questions:
                content = q.content or {}
                if content.get("status") == "answered":
                    answered.append({
                        "code": q.artifact_code,
                        "question": content.get("question"),
                        "answer": content.get("answer"),
                        "related": q.parent_refs or []
                    })
            return answered
        except Exception as e:
            logger.warning(f"Could not get answered questions: {e}")
            return []
    
    def _store_questions(self, questions: List[Dict], execution_id: int) -> None:
        """Store architect questions as artifacts"""
        try:
            from app.models.artifact import ExecutionArtifact
            for q in questions:
                existing = self.db.query(ExecutionArtifact).filter(
                    ExecutionArtifact.execution_id == execution_id,
                    ExecutionArtifact.artifact_code == q.get("artifact_code")
                ).first()
                
                if not existing:
                    artifact = ExecutionArtifact(
                        execution_id=execution_id,
                        artifact_type="question",
                        artifact_code=q.get("artifact_code"),
                        title=q.get("title", "Question from Architect"),
                        producer_agent="architect",
                        content={
                            "question": q.get("content", {}).get("question", q.get("question", "")),
                            "context": q.get("content", {}).get("context", ""),
                            "related_artifacts": q.get("parent_refs", []),
                            "status": "pending"
                        },
                        parent_refs=q.get("parent_refs", []),
                        status="active"
                    )
                    self.db.add(artifact)
            self.db.commit()
        except Exception as e:
            logger.error(f"Failed to store questions: {e}")
    
    async def _ba_answer_questions(
        self,
        project: "Project",
        questions: List[Dict],
        ba_artifacts: Dict,
        execution: "Execution",
        agent_status: Dict
    ) -> Dict[str, Any]:
        """Have BA answer architect's questions"""
        logger.info(f"[V2 Iteration] BA answering {len(questions)} questions")
        
        # Build answer prompt
        answer_prompt = self._build_answer_prompt(questions, ba_artifacts)
        
        # Temporarily update BA status
        agent_status["state"] = "running"
        agent_status["message"] = f"Answering {len(questions)} clarification questions..."
        self.db.commit()
        
        try:
            # Use OpenAI directly for quick answer
            import openai
            client = openai.OpenAI()
            
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "You are a Salesforce Business Analyst. Answer the architect's questions based on the business requirements context."},
                    {"role": "user", "content": answer_prompt}
                ],
                max_tokens=4000,
                temperature=0.3
            )
            
            answer_text = response.choices[0].message.content
            
            return {
                "success": True,
                "answers": answer_text,
                "questions_answered": len(questions)
            }
            
        except Exception as e:
            logger.error(f"BA answer failed: {e}")
            return {"success": False, "error": str(e)}
    
    def _build_answer_prompt(self, questions: List[Dict], ba_artifacts: Dict) -> str:
        """Build prompt for BA to answer questions"""
        lines = ["# QUESTIONS FROM SOLUTION ARCHITECT", ""]
        lines.append("Please answer these clarification questions about the business requirements:")
        lines.append("")
        
        for q in questions:
            code = q.get("artifact_code", "Q-???")
            content = q.get("content", {})
            question_text = content.get("question", q.get("question", "No question text"))
            context = content.get("context", "")
            related = q.get("parent_refs", [])
            
            lines.append(f"## {code}")
            if context:
                lines.append(f"**Context:** {context}")
            lines.append(f"**Question:** {question_text}")
            if related:
                lines.append(f"**Related to:** {', '.join(related)}")
            lines.append("")
        
        lines.append("---")
        lines.append("For each question, provide:")
        lines.append("1. A clear, specific answer")
        lines.append("2. Any additional business rules or constraints")
        lines.append("3. Recommendation for the architect if applicable")
        
        return "\n".join(lines)
    
    def _update_questions_with_answers(self, ba_result: Dict, execution_id: int) -> None:
        """Update question artifacts with BA answers"""
        try:
            from app.models.artifact import ExecutionArtifact
            answers_text = ba_result.get("answers", "")
            
            # Mark all pending questions as answered
            questions = self.db.query(ExecutionArtifact).filter(
                ExecutionArtifact.execution_id == execution_id,
                ExecutionArtifact.artifact_type == "question"
            ).all()
            
            for q in questions:
                content = q.content or {}
                if content.get("status") == "pending":
                    content["status"] = "answered"
                    content["answer"] = f"See BA response in iteration"
                    content["full_response"] = answers_text
                    q.content = content
            
            self.db.commit()
        except Exception as e:
            logger.error(f"Failed to update questions: {e}")


    async def _execute_single_agent(
        self,
        agent_id: str,
        project: Project,
        previous_outputs: Dict[str, Any],
        execution: Execution,
        agent_status: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        Execute a single agent using the AgentIntegrationService
        """

        # Prepare project data for agent
        project_data = {
            "name": project.name,
            "salesforce_product": project.salesforce_product,
            "organization_type": project.organization_type,
            "business_requirements": project.business_requirements,
            "existing_systems": project.existing_systems,
            "compliance_requirements": project.compliance_requirements,
            "expected_users": project.expected_users,
            "expected_data_volume": project.expected_data_volume,
            "architecture_preferences": project.architecture_preferences or {}
        }

        # Update progress: Starting execution
        agent_status["progress"] = 10
        agent_status["message"] = "Initializing agent..."
        execution.agent_execution_status[agent_id] = agent_status
        self.db.commit()

        try:
            # Execute agent via AgentIntegrationService
            logger.info(f"Executing agent {agent_id} for project {project.id}")

            agent_status["progress"] = 20
            agent_status["message"] = "Executing agent..."
            execution.agent_execution_status[agent_id] = agent_status
            self.db.commit()

            result = await self.agent_service.execute_agent(
                agent_id=agent_id,
                project_data=project_data,
                execution_id=execution.id,
                previous_outputs=previous_outputs,
                timeout=300  # 5 minutes per agent
            )

            # Update progress: Agent completed
            agent_status["progress"] = 80
            agent_status["message"] = "Validating output quality..."
            execution.agent_execution_status[agent_id] = agent_status
            self.db.commit()

            # Validate output quality
            quality_results = await self.quality_service.validate_agent_output(
                agent_id=agent_id,
                output=result.get("output", {}),
                previous_outputs=previous_outputs
            )

            quality_summary = self.quality_service.aggregate_results(quality_results)

            # Log quality check results
            logger.info(
                f"Agent {agent_id} quality check: "
                f"{quality_summary['overall_status']} "
                f"({quality_summary['passed']}/{quality_summary['total_gates']} passed)"
            )

            # Add quality info to result
            result["quality_check"] = quality_summary

            # If quality gates failed critically, log warning but continue
            if quality_summary["overall_status"] == QualityGateStatus.FAILED.value:
                logger.warning(f"Agent {agent_id} failed quality gates but continuing execution")

            return result

        except AgentExecutionError as e:
            logger.error(f"Agent {agent_id} execution failed: {str(e)}")
            # Return error result
            return {
                "agent_id": agent_id,
                "status": "failed",
                "output": None,
                "error": str(e),
                "execution_time": 0
            }

    def _generate_agent_output(
        self,
        agent_id: str,
        project: Project,
        context: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Generate agent-specific output content"""

        agent_outputs = {
            'ba': {
                "title": "Business Requirements Document",
                "content": f"""# Business Analysis for {project.name}

## Project Overview
This {project.salesforce_product} implementation aims to transform {project.organization_type.lower()} business processes through strategic Salesforce deployment.

## Business Objectives
{project.business_requirements}

## Stakeholder Analysis
Key stakeholders have been identified across various departments. Their requirements have been documented and prioritized according to business impact.

## Success Criteria
- User adoption rate > 80%
- Process efficiency improvement > 30%
- Data accuracy improvement > 95%

## Risk Assessment
Potential risks have been identified and mitigation strategies prepared.""",
                "deliverables": [
                    "Business Requirements Document (BRD)",
                    "Stakeholder Analysis Report",
                    "Success Metrics Definition"
                ]
            },
            'architect': {
                "title": "Solution Architecture Document",
                "content": f"""# Solution Architecture for {project.name}

## Architecture Overview
Proposed architecture leverages {project.salesforce_product} capabilities to create a scalable, maintainable solution.

## System Components
- Salesforce Core Platform
- Custom Objects and Fields
- Integration Layer
- External System Connectors

## Data Model
Custom objects designed to support business processes while maintaining data integrity and performance.

## Integration Architecture
{f'Integration with: {project.existing_systems}' if project.existing_systems else 'Standalone implementation with future integration capabilities.'}

## Security & Compliance
{project.compliance_requirements if project.compliance_requirements else 'Standard Salesforce security model with role-based access control.'}

## Scalability Considerations
Architecture designed to support {project.expected_users or 'growing number of'} users with optimal performance.""",
                "deliverables": [
                    "Solution Architecture Document",
                    "Data Model Diagram",
                    "Integration Architecture Diagram"
                ]
            },
            'apex': {
                "title": "Apex Development Specifications",
                "content": f"""# Apex Development Specifications for {project.name}

## Apex Classes Overview
Custom business logic implementation following Salesforce best practices.

## Trigger Framework
- Bulkified trigger handlers
- Trigger control mechanism
- Recursive trigger prevention

## Service Layer
- Business logic encapsulation
- Reusable service classes
- Error handling framework

## Integration Classes
- REST/SOAP API integrations
- Callout handling
- Response parsing

## Code Quality Standards
- Minimum 85% code coverage
- Governor limit optimization
- Comprehensive error handling""",
                "deliverables": [
                    "Apex Class Specifications",
                    "Trigger Framework Design",
                    "Integration Class Documentation"
                ]
            },
            'lwc': {
                "title": "Lightning Web Components Specifications",
                "content": f"""# Lightning Web Components for {project.name}

## Component Architecture
Modern, performant LWC components following Salesforce best practices.

## Component Library
- Data display components
- Input form components
- Navigation components
- Integration components

## Design System
Leveraging Salesforce Lightning Design System (SLDS) for consistent UI/UX.

## Performance Optimization
- Lazy loading
- Client-side caching
- Efficient data binding

## Accessibility
WCAG 2.1 AA compliance ensuring inclusive user experience.""",
                "deliverables": [
                    "LWC Component Specifications",
                    "Component Architecture Diagram",
                    "User Interface Mockups"
                ]
            },
            'admin': {
                "title": "Configuration & Administration Guide",
                "content": f"""# Configuration Specifications for {project.name}

## Declarative Features
Maximizing declarative configuration before custom code.

## Objects & Fields
- Custom objects configuration
- Field-level security
- Validation rules
- Record types and page layouts

## Automation
- Process Builder flows
- Flow Builder automations
- Workflow rules (legacy migration)
- Approval processes

## Security Configuration
- Profiles and permission sets
- Sharing rules
- Field-level security
- Object permissions

## User Management
Support for {project.expected_users or 'enterprise-scale'} user base.""",
                "deliverables": [
                    "Configuration Guide",
                    "Automation Specifications",
                    "Security Matrix"
                ]
            },
            'qa': {
                "title": "Quality Assurance & Testing Strategy",
                "content": f"""# QA Strategy for {project.name}

## Testing Approach
Comprehensive testing strategy ensuring quality and reliability.

## Test Levels
- Unit Testing (Apex, LWC)
- Integration Testing
- System Testing
- User Acceptance Testing (UAT)

## Test Coverage
- Minimum 85% code coverage
- All critical paths tested
- Edge cases validated

## Test Automation
- Automated regression suite
- Continuous integration testing
- Performance testing

## Quality Metrics
- Defect density targets
- Test coverage metrics
- Performance benchmarks""",
                "deliverables": [
                    "Test Strategy Document",
                    "Test Cases Repository",
                    "UAT Plan"
                ]
            },
            'devops': {
                "title": "DevOps & Deployment Strategy",
                "content": f"""# DevOps Strategy for {project.name}

## Deployment Model
Structured deployment pipeline ensuring reliable releases.

## Environment Strategy
- Development (Dev)
- System Integration Testing (SIT)
- User Acceptance Testing (UAT)
- Production

## CI/CD Pipeline
- Automated builds
- Automated testing
- Deployment automation
- Rollback procedures

## Version Control
Git-based version control with branching strategy.

## Monitoring & Maintenance
- Production monitoring
- Performance tracking
- Issue management""",
                "deliverables": [
                    "Deployment Strategy Document",
                    "CI/CD Pipeline Configuration",
                    "Environment Management Plan"
                ]
            },
            'data': {
                "title": "Data Migration Strategy",
                "content": f"""# Data Migration Strategy for {project.name}

## Migration Approach
Structured data migration ensuring data integrity and minimal downtime.

## Data Sources
{f'Migrating from: {project.existing_systems}' if project.existing_systems else 'New implementation - initial data load.'}

## Migration Process
- Data extraction
- Data transformation and cleansing
- Data validation
- Data loading
- Post-migration verification

## Data Volume
Expected volume: {project.expected_data_volume or 'Standard enterprise volume'}

## Data Quality
- Deduplication
- Validation rules
- Data enrichment
- Quality metrics""",
                "deliverables": [
                    "Data Migration Plan",
                    "Data Mapping Document",
                    "Migration Scripts"
                ]
            },
            'trainer': {
                "title": "Training & Documentation Plan",
                "content": f"""# Training Strategy for {project.name}

## Training Approach
Comprehensive training ensuring user adoption and proficiency.

## Training Modules
- End-user training
- Administrator training
- Developer training
- Executive overview

## Training Delivery
- Instructor-led training sessions
- E-learning modules
- Quick reference guides
- Video tutorials

## Documentation
- User guides
- Administrator guides
- Technical documentation
- FAQs and troubleshooting

## User Adoption
Target: {project.expected_users or 'All'} users trained and certified.""",
                "deliverables": [
                    "Training Plan",
                    "Training Materials",
                    "User Documentation"
                ]
            },
            'pm': {
                "title": "Project Management & Integration Summary",
                "content": f"""# Project Summary for {project.name}

## Executive Summary
This comprehensive Solution Design Specification documents the complete technical approach for implementing {project.salesforce_product} for {project.organization_type.lower()}.

## Project Scope
The solution addresses key business requirements while maintaining technical excellence and adherence to best practices.

## Team Deliverables
All specialized teams have completed their analysis and design work, producing comprehensive specifications ready for implementation.

## Timeline & Milestones
- Requirements Analysis: ✓ Complete
- Solution Design: ✓ Complete
- Technical Specifications: ✓ Complete
- Ready for Implementation

## Success Criteria
Project success measured against defined business objectives with clear KPIs and success metrics.

## Risk Management
All identified risks have mitigation strategies in place.

## Next Steps
Implementation phase ready to commence with clear technical specifications and comprehensive documentation.""",
                "deliverables": [
                    "Complete SDS Document",
                    "Project Plan",
                    "Risk Register",
                    "Success Metrics Dashboard"
                ]
            }
        }

        return {
            "agent_id": agent_id,
            **agent_outputs.get(agent_id, {
                "title": f"{agent_id.upper()} Specifications",
                "content": f"Specifications for {agent_id}",
                "deliverables": [f"{agent_id.upper()} Documentation"]
            }),
            "timestamp": datetime.now(timezone.utc).isoformat()
        }

    async def _generate_sds_document(
        self,
        project: Project,
        agent_outputs: Dict[str, Any],
        execution_id: int
    ) -> str:
        """
        Generate professional SDS document using template-based generator
        
        Uses Digital-Humans.fr template with 8 sections:
        1. Introduction
        2. Solution Scope
        3. Roles & Profiles
        4. Functional Design (per module)
        5. Data Migration
        6. Systems Interface
        7. Assumptions & Dependencies
        8. Signoff
        """
        try:
            # Use new template-based generator (reads from DB)
            output_path = generate_sds_from_template(execution_id)
            logger.info(f"Template SDS document generated: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error with template generator, using fallback: {e}")
            import traceback
            traceback.print_exc()
            # Continue with basic generation below
        
        # ===== FALLBACK: Basic document generation =====
        doc = Document()

        # ===== TITLE PAGE =====
        title = doc.add_heading("Solution Design Specification", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Project name
        project_para = doc.add_paragraph(project.name)
        project_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        project_para.runs[0].font.size = Pt(16)
        project_para.runs[0].font.bold = True

        # Generated by - IMPORTANT: "Digital Humans System"
        generated_para = doc.add_paragraph("Generated by Digital Humans System")
        generated_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        generated_para.runs[0].font.size = Pt(10)
        generated_para.runs[0].font.color.rgb = RGBColor(107, 114, 128)

        # Date
        date_para = doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.runs[0].font.size = Pt(10)

        # Product info
        product_para = doc.add_paragraph(
            f"{project.salesforce_product} • {project.organization_type}"
        )
        product_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        product_para.runs[0].font.size = Pt(12)
        product_para.runs[0].font.color.rgb = RGBColor(37, 99, 235)

        doc.add_page_break()

        # ===== EXECUTIVE SUMMARY =====
        doc.add_heading("Executive Summary", 1)

        summary = f"""This Solution Design Specification outlines the technical implementation for {project.name}, a {project.salesforce_product} solution for {project.organization_type.lower()}.

The Digital Humans AI system has orchestrated {len(agent_outputs)} specialized AI agents to produce comprehensive technical specifications. Each agent contributed their domain expertise to create a complete, implementable solution design.

This document provides detailed specifications across all aspects of the implementation, from business analysis through deployment strategy."""

        doc.add_paragraph(summary)

        # ===== BUSINESS REQUIREMENTS =====
        doc.add_heading("Business Requirements", 1)

        # CRITICAL: Limit to 3-7 lines
        requirements = project.business_requirements or "No requirements specified"
        req_lines = [line.strip() for line in requirements.strip().split('\n') if line.strip()]

        if len(req_lines) > 7:
            req_lines = req_lines[:7]

        for line in req_lines:
            p = doc.add_paragraph(line, style='List Bullet')
            p.paragraph_format.space_after = Pt(6)

        # ===== TECHNICAL CONTEXT =====
        if project.existing_systems or project.compliance_requirements or project.expected_users:
            doc.add_page_break()
            doc.add_heading("Technical Context", 1)

            if project.existing_systems:
                doc.add_heading("Existing Systems & Integrations", 2)
                doc.add_paragraph(project.existing_systems)

            if project.compliance_requirements:
                doc.add_heading("Compliance & Security Requirements", 2)
                doc.add_paragraph(project.compliance_requirements)

            if project.expected_users:
                doc.add_heading("Scale & Performance", 2)
                scale_text = f"Expected users: {project.expected_users}"
                if project.expected_data_volume:
                    scale_text += f"\nExpected data volume: {project.expected_data_volume}"
                doc.add_paragraph(scale_text)

        # ===== AGENT OUTPUTS =====
        agent_names = {
            'ba': 'Business Analysis',
            'architect': 'Solution Architecture',
            'apex': 'Apex Development Specifications',
            'lwc': 'Lightning Web Components',
            'admin': 'Configuration & Administration',
            'qa': 'Quality Assurance & Testing',
            'devops': 'DevOps & Deployment',
            'data': 'Data Migration',
            'trainer': 'Training & Documentation',
            'pm': 'Project Management Summary'
        }

        for agent_id, result in agent_outputs.items():
            doc.add_page_break()

            # Section heading
            section_name = agent_names.get(agent_id, agent_id.upper())
            doc.add_heading(section_name, 1)

            # Extract output from result (agent integration service wraps output)
            # Structure: result["output"]["json_data"]["content"]["raw_markdown"]
            output = result.get("output", result)
            
            # Handle nested structure from AgentIntegrationService
            json_data = output.get("json_data", output) if isinstance(output, dict) else output
            
            # Get content - could be in json_data["content"]["raw_markdown"] or json_data["content"]
            content = "No content generated"
            if isinstance(json_data, dict):
                content_obj = json_data.get("content", {})
                if isinstance(content_obj, dict):
                    content = content_obj.get("raw_markdown", content_obj.get("content", "No content generated"))
                elif isinstance(content_obj, str):
                    content = content_obj
                    
            logger.info(f"Agent {agent_id} content length: {len(content)} chars")

            # Parse markdown-like content
            for line in content.split('\n'):
                line = line.strip()
                if not line:
                    continue

                if line.startswith('# '):
                    doc.add_heading(line[2:], 2)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], 3)
                elif line.startswith('- '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                else:
                    doc.add_paragraph(line)

            # Deliverables
            if 'deliverables' in output and output['deliverables']:
                doc.add_heading("Key Deliverables", 3)
                for deliverable in output['deliverables']:
                    doc.add_paragraph(deliverable, style='List Bullet')

        # ===== SAVE DOCUMENT =====
        output_dir = "/home/user/front-end-digital-humans/outputs"
        os.makedirs(output_dir, exist_ok=True)

        safe_name = "".join(c for c in project.name if c.isalnum() or c in (' ', '-', '_')).rstrip()
        safe_name = safe_name.replace(' ', '_')
        output_path = f"{output_dir}/SDS_{execution_id}_{safe_name}.docx"

        doc.save(output_path)

        return output_path



    # ============ V2 ARTIFACT METHODS ============
    
    def _initialize_validation_gates(self, execution_id: int) -> List[Dict]:
        """Initialize the 6 validation gates for an execution"""
        gates = self.artifact_service.initialize_gates(execution_id)
        logger.info(f"Initialized {len(gates)} validation gates for execution {execution_id}")
        return [g.to_dict() for g in gates]
    
    def _get_artifact_type_for_agent(self, agent_id: str) -> Optional[str]:
        """Map agent ID to primary artifact type"""
        mapping = {
            'ba': 'business_req',      # BA produces BR and UC
            'architect': 'spec',        # Architect produces ADR and SPEC
            'apex': 'code',             # Apex produces CODE
            'lwc': 'code',              # LWC produces CODE
            'admin': 'config',          # Admin produces CONFIG
            'qa': 'test',               # QA produces TEST
            'devops': 'config',         # DevOps produces CONFIG
            'data': 'config',           # Data produces CONFIG
            'trainer': 'doc',           # Trainer produces DOC
        }
        return mapping.get(agent_id)
    
    def _extract_artifacts_from_agent_output(
        self,
        agent_id: str,
        output: Dict[str, Any],
        execution_id: int
    ) -> List[Dict]:
        """
        Extract structured artifacts from agent output.
        Parses the markdown content for BR-xxx, UC-xxx, ADR-xxx, SPEC-xxx, CODE-xxx, CFG-xxx, TEST-xxx, DOC-xxx patterns.
        """
        artifacts = []
        
        try:
            # Get the content from output - handle nested structure
            json_data = output.get("output", {})
            if isinstance(json_data, dict):
                json_data = json_data.get("json_data", json_data)
            
            if not isinstance(json_data, dict):
                json_data = {}
            
            # Get raw markdown content - handle different structures
            content_field = json_data.get("content", {})
            if isinstance(content_field, dict):
                raw_markdown = content_field.get("raw_markdown", "")
            elif isinstance(content_field, str):
                raw_markdown = content_field
            else:
                raw_markdown = str(content_field) if content_field else ""
            
            if not raw_markdown:
                logger.warning(f"No raw_markdown found in output for agent {agent_id}")
                return artifacts
            
            # Define which patterns to look for based on agent
            agent_pattern_map = {
                'ba': [('BR', 'business_req'), ('UC', 'use_case')],
                'architect': [('ADR', 'adr'), ('SPEC', 'spec'), ('Q', 'question')],
                'apex': [('CODE', 'code')],
                'lwc': [('CODE', 'code')],
                'admin': [('CFG', 'config')],
                'qa': [('TEST', 'test')],
                'trainer': [('DOC', 'doc')],
            }
            
            patterns_to_check = agent_pattern_map.get(agent_id, [])
            
            for prefix, artifact_type in patterns_to_check:
                # Pattern to match "### BR-001: Title" and capture content until next artifact or ---
                pattern = rf'###\s+{prefix}-(\d+):\s*([^\n]+)\n(.+?)(?=###\s+(?:BR|UC|ADR|SPEC|CODE|CFG|TEST|DOC|Q)-\d+:|---\s*\n|\Z)'
                
                matches = re.finditer(pattern, raw_markdown, re.DOTALL | re.IGNORECASE)
                
                for match in matches:
                    artifact_num = match.group(1)
                    title = match.group(2).strip()
                    artifact_content = match.group(3).strip()
                    
                    # Build artifact code
                    artifact_code = f"{prefix}-{artifact_num.zfill(3)}"
                    
                    # Parse structured fields from content
                    parsed_content = self._parse_artifact_fields(artifact_content)
                    parsed_content["raw_content"] = artifact_content[:10000]
                    
                    # Find parent references
                    parent_refs = self._extract_parent_refs(artifact_content)
                    
                    artifact_data = {
                        "execution_id": execution_id,
                        "artifact_type": artifact_type,
                        "artifact_code": artifact_code,
                        "title": title[:255],
                        "producer_agent": agent_id,
                        "content": parsed_content,
                        "parent_refs": parent_refs
                    }
                    
                    artifacts.append(artifact_data)
                    logger.info(f"Extracted artifact {artifact_code}: {title[:50]}...")
            
            # If no structured artifacts found, create a single artifact for the whole output
            if not artifacts:
                logger.info(f"No structured artifacts found for {agent_id}, creating single artifact")
                artifact_type = self._get_artifact_type_for_agent(agent_id)
                if artifact_type:
                    prefix_map = {
                        'business_req': 'BR',
                        'use_case': 'UC',
                        'spec': 'SPEC',
                        'adr': 'ADR',
                        'code': 'CODE',
                        'config': 'CFG',
                        'test': 'TEST',
                        'doc': 'DOC'
                    }
                    prefix = prefix_map.get(artifact_type, 'ART')
                    
                    artifacts.append({
                        "execution_id": execution_id,
                        "artifact_type": artifact_type,
                        "artifact_code": f"{prefix}-001",
                        "title": f"{agent_id.upper()} Output",
                        "producer_agent": agent_id,
                        "content": {
                            "raw_content": raw_markdown[:50000],
                            "metadata": json_data.get("metadata", {})
                        },
                        "parent_refs": []
                    })
            
            logger.info(f"Extracted {len(artifacts)} artifacts from {agent_id}")
            return artifacts
            
        except Exception as e:
            logger.error(f"Error extracting artifacts from {agent_id}: {str(e)}")
            import traceback
            logger.error(traceback.format_exc())
            return artifacts

    def _parse_artifact_fields(self, content: str) -> Dict[str, Any]:
        """Parse structured fields from artifact content."""
        parsed = {}
        
        # Common field patterns
        field_patterns = {
            "priority": r'\*\*Priority:\*\*\s*(.+)',
            "category": r'\*\*Category:\*\*\s*(.+)',
            "description": r'\*\*Description:\*\*\s*(.+?)(?=\*\*|$)',
            "parent_br": r'\*\*Parent BR:\*\*\s*(BR-\d+)',
            "related_uc": r'\*\*Related UC:\*\*\s*([UC\d,\s-]+)',
            "related_spec": r'\*\*Related SPEC:\*\*\s*(SPEC-\d+)',
            "related_adr": r'\*\*Related ADR:\*\*\s*(ADR-\d+)',
            "type": r'\*\*Type:\*\*\s*(.+)',
            "actor": r'\*\*Actor:\*\*\s*(.+)',
            "status": r'\*\*Status:\*\*\s*(.+)',
            "assigned_to": r'\*\*Assigned To:\*\*\s*(.+)',
            "stakeholder": r'\*\*Stakeholder:\*\*\s*(.+)',
            "business_value": r'\*\*Business Value:\*\*\s*(.+?)(?=\*\*|$)',
            "salesforce_feature": r'\*\*Salesforce Feature:\*\*\s*(.+)',
        }
        
        for field, pattern in field_patterns.items():
            match = re.search(pattern, content, re.IGNORECASE | re.DOTALL)
            if match:
                value = match.group(1).strip()
                # Clean up the value (remove trailing newlines, limit length)
                value = value.split('\n')[0].strip()[:500]
                parsed[field] = value
        
        # Extract acceptance criteria
        ac_match = re.search(r'\*\*Acceptance Criteria:\*\*\s*(.+?)(?=\*\*[A-Z]|---\s*\n|\Z)', content, re.DOTALL | re.IGNORECASE)
        if ac_match:
            criteria = re.findall(r'- \[[ x]\]\s*(.+)', ac_match.group(1))
            if criteria:
                parsed["acceptance_criteria"] = [c.strip() for c in criteria[:20]]
        
        return parsed

    def _extract_parent_refs(self, content: str) -> List[str]:
        """Extract references to parent artifacts (BR-xxx, UC-xxx, etc.)."""
        refs = []
        
        # Look for explicit parent references
        parent_patterns = [
            r'\*\*Parent BR:\*\*\s*(BR-\d+)',
            r'\*\*Related UC:\*\*\s*((?:UC-\d+[,\s]*)+)',
            r'\*\*Related SPEC:\*\*\s*(SPEC-\d+)',
            r'\*\*Related ADR:\*\*\s*(ADR-\d+)',
            r'\*\*Related CODE:\*\*\s*((?:CODE-\d+[,\s]*)+)',
        ]
        
        for pattern in parent_patterns:
            matches = re.findall(pattern, content, re.IGNORECASE)
            for match in matches:
                # Split if multiple refs (e.g., "UC-001, UC-002")
                individual_refs = re.findall(r'(?:BR|UC|ADR|SPEC|CODE|CFG|TEST|DOC)-\d+', match)
                refs.extend(individual_refs)
        
        return list(set(refs))[:10]  # Deduplicate and limit


    def _persist_artifacts(self, artifacts: List[Dict], execution_id: int) -> List[int]:
        """Persist extracted artifacts to database"""
        created_ids = []
        
        for artifact_data in artifacts:
            try:
                # Map string type to enum
                type_str = artifact_data.get("artifact_type", "doc")
                type_enum = ArtifactTypeEnum(type_str)
                
                artifact = self.artifact_service.create_artifact(ArtifactCreate(
                    execution_id=execution_id,
                    artifact_type=type_enum,
                    artifact_code=artifact_data["artifact_code"],
                    title=artifact_data["title"],
                    producer_agent=artifact_data["producer_agent"],
                    content=artifact_data["content"],
                    parent_refs=artifact_data.get("parent_refs", [])
                ))
                created_ids.append(artifact.id)
                logger.info(f"Created artifact {artifact.artifact_code} (ID: {artifact.id})")
                
            except Exception as e:
                logger.error(f"Error persisting artifact {artifact_data.get('artifact_code')}: {e}")
        
        return created_ids
    
    def _check_gate_progress(self, execution_id: int, agent_id: str) -> Optional[Dict]:
        """Check if a validation gate should be triggered after this agent"""
        # Gate 1: After BA (BR/UC validation)
        if agent_id == "ba":
            return {
                "gate_number": 1,
                "gate_name": "BR/UC Validation",
                "artifact_types": ["business_req", "use_case"]
            }
        
        # Gate 2: After Architect (SDS/ADR validation)
        if agent_id == "architect":
            return {
                "gate_number": 2,
                "gate_name": "SDS/ADR Validation",
                "artifact_types": ["adr", "spec"]
            }
        
        # Gate 3: After all developers (Code validation)
        if agent_id in ["apex", "lwc", "admin"]:
            return {
                "gate_number": 3,
                "gate_name": "Code/Config Validation",
                "artifact_types": ["code", "config"]
            }
        
        # Gate 4: After QA
        if agent_id == "qa":
            return {
                "gate_number": 4,
                "gate_name": "QA Validation",
                "artifact_types": ["test"]
            }
        
        # Gate 5: After Trainer
        if agent_id == "trainer":
            return {
                "gate_number": 5,
                "gate_name": "Training Validation",
                "artifact_types": ["doc"]
            }
        
        return None
    
    def _update_gate_artifacts_count(self, execution_id: int, gate_number: int) -> None:
        """Update the artifacts count for a gate"""
        try:
            gate = self.artifact_service.get_gate(execution_id, gate_number)
            if gate:
                artifact_types = gate.artifact_types or []
                total = 0
                for art_type in artifact_types:
                    artifacts = self.artifact_service.list_artifacts(
                        execution_id, artifact_type=art_type
                    )
                    total += len(artifacts)
                
                gate.artifacts_count = total
                self.db.commit()
                logger.info(f"Gate {gate_number} artifacts count updated to {total}")
        except Exception as e:
            logger.error(f"Error updating gate artifacts count: {e}")

# Background task wrapper
async def execute_agents_background(
    execution_id: int,
    project_id: int,
    selected_agents: List[str]
):
    """
    Wrapper for background execution
    This is called from the API route
    """
    db = SessionLocal()
    try:
        service = PMOrchestratorService(db)
        await service.execute_agents(
            execution_id=execution_id,
            project_id=project_id,
            selected_agents=selected_agents
        )
    except Exception as e:
        print(f"Background execution error: {str(e)}")
        import traceback
        traceback.print_exc()
        raise
    finally:
        db.close()
