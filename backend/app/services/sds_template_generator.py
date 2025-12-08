#!/usr/bin/env python3
"""
Professional SDS Document Generator - Template Based
Based on Bacon Formula / Digital-Humans.fr template structure

Generates complete Solution Design Specification documents
following the standard 8-section structure.
"""
import os
import json
import re
from datetime import datetime
from typing import Dict, List, Any, Optional
from collections import defaultdict

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import parse_xml

from sqlalchemy import create_engine, text

# Configuration
DATABASE_URL = os.getenv("DATABASE_URL", "postgresql://digital_humans:DH_SecurePass2025!@127.0.0.1:5432/digital_humans_db")
OUTPUT_DIR = "/app/outputs"
COMPANY_NAME = "Digital-Humans.fr"
COMPANY_ADDRESS = "Paris, France"

# Colors
HEADER_BLUE = RGBColor(31, 78, 121)      # #1F4E79
LIGHT_BLUE = RGBColor(68, 114, 196)      # #4472C4
GRAY_TEXT = RGBColor(107, 114, 128)      # #6B7280
WHITE = RGBColor(255, 255, 255)
BLACK = RGBColor(0, 0, 0)


def set_cell_shading(cell, color_hex: str):
    """Set cell background color"""
    shading = parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_table_header_row(row, color_hex: str = "1F4E79"):
    """Format a table header row with background color and white bold text"""
    for cell in row.cells:
        set_cell_shading(cell, color_hex)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True


def create_table_with_header(doc, headers: List[str], data: List[List[str]], col_widths: List[float] = None) -> None:
    """Create a formatted table with header row"""
    if not data:
        doc.add_paragraph("No data available.")
        return
    
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Header row
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
    set_table_header_row(table.rows[0])
    
    # Data rows
    for row_data in data:
        row = table.add_row().cells
        for i, value in enumerate(row_data):
            row[i].text = str(value) if value else ""
    
    # Set column widths if provided
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)
    
    doc.add_paragraph()


class SDSTemplateGenerator:
    """
    Professional SDS Document Generator following template structure
    """
    
    def __init__(self, execution_id: int):
        self.execution_id = execution_id
        self.doc = None
        self.data = {}
        self.project_info = {}
        self.modules = []
        
    def load_data(self):
        """Load all data from database"""
        engine = create_engine(DATABASE_URL)
        
        with engine.connect() as conn:
            # Get execution and project info
            exec_info = conn.execute(text("""
                SELECT e.id, e.project_id, p.name as project_name,
                       p.salesforce_product, p.organization_type, e.created_at
                FROM executions e
                JOIN projects p ON e.project_id = p.id
                WHERE e.id = :exec_id
            """), {"exec_id": self.execution_id}).fetchone()
            
            if exec_info:
                self.project_info = {
                    "execution_id": exec_info[0],
                    "project_id": exec_info[1],
                    "project_name": exec_info[2] or "Salesforce Implementation",
                    "client_name": "Client",  # Derived from project name
                    "salesforce_product": exec_info[3] or "Sales Cloud",
                    "org_type": exec_info[4] or "New Implementation",
                    "date": exec_info[5].strftime("%Y-%m-%d") if exec_info[5] else datetime.now().strftime("%Y-%m-%d")
                }
            else:
                self.project_info = {
                    "execution_id": self.execution_id,
                    "project_id": 0,
                    "project_name": "Salesforce Implementation",
                    "client_name": "Client",
                    "salesforce_product": "Sales Cloud",
                    "org_type": "New Implementation",
                    "date": datetime.now().strftime("%Y-%m-%d")
                }
            
            # Get artifacts
            artifacts = conn.execute(text("""
                SELECT producer_agent, content
                FROM execution_artifacts
                WHERE execution_id = :exec_id AND artifact_type = 'output'
            """), {"exec_id": self.execution_id}).fetchall()
            
            for agent, content in artifacts:
                self.data[agent] = content if isinstance(content, dict) else json.loads(content)
        
        # Process modules from use cases
        self._process_modules()
        
        print(f"Loaded data: {list(self.data.keys())}")
        print(f"Project: {self.project_info['project_name']}")
        print(f"Modules identified: {len(self.modules)}")
    
    def _process_modules(self):
        """Group use cases into functional modules"""
        ba_data = self.data.get("olivia", {})
        use_cases = ba_data.get("use_cases", [])
        
        # Module definitions based on UC prefixes
        module_defs = {
            "001": {"name": "Lead Capture", "purpose": "Multi-channel lead acquisition and initial processing"},
            "002": {"name": "Lead Qualification", "purpose": "Lead scoring, qualification stages, and conversion"},
            "003": {"name": "Territory Management", "purpose": "Geographic territory assignment and routing"},
            "004": {"name": "Sales Rep Assignment", "purpose": "Availability-based lead distribution"},
            "005": {"name": "Behavioral Scoring", "purpose": "Engagement tracking and lead prioritization"},
            "006": {"name": "Opportunity Management", "purpose": "Multi-product opportunity and deal management"},
            "007": {"name": "Trade-In Management", "purpose": "Vehicle trade-in requests and appraisals"},
            "009": {"name": "Approval Workflows", "purpose": "Multi-level approval for discounts and offers"},
            "010": {"name": "Financing Configuration", "purpose": "Loan and financing product configuration"},
            "011": {"name": "Insurance Products", "purpose": "Insurance product recommendations and sales"},
            "012": {"name": "Complementary Services", "purpose": "Add-on services and bundles"},
            "013": {"name": "Dealership Reporting", "purpose": "Multi-dealership consolidated reporting"},
            "014": {"name": "Brand Analytics", "purpose": "Performance analysis by vehicle brand"},
            "015": {"name": "Vehicle Reporting", "purpose": "Individual vehicle performance tracking"},
            "016": {"name": "Customer Segmentation", "purpose": "CRM Analytics and customer segments"},
            "017": {"name": "Lead Source Analytics", "purpose": "Lead source ROI and attribution"},
            "018": {"name": "Territory Planning", "purpose": "Territory hierarchy and planning"},
            "019": {"name": "Territory Exceptions", "purpose": "Manual territory overrides"},
            "024": {"name": "Discount Approvals", "purpose": "Discount threshold approval workflows"},
            "025": {"name": "Destocking Offers", "purpose": "Inventory destocking approval process"},
            "026": {"name": "Custom Offers", "purpose": "Custom offer approval workflow"},
            "027": {"name": "Sales Forecasting", "purpose": "Pipeline and revenue forecasting"},
            "028": {"name": "Forecast Adjustments", "purpose": "Historical-based forecast modifications"},
            "029": {"name": "Vehicle Valuation", "purpose": "External valuation service integration"},
            "030": {"name": "Valuation Sync", "purpose": "Trade-in status synchronization"},
            "031": {"name": "Shopping Cart", "purpose": "Offer management and cart operations"},
            "032": {"name": "Margin Calculation", "purpose": "Dynamic margin and profitability"},
            "033": {"name": "Notifications", "purpose": "Automated alerts and notifications"},
            "034": {"name": "Quote Generation", "purpose": "PDF quote creation and delivery"},
            "035": {"name": "Document Branding", "purpose": "Dealership-branded documents"},
            "036": {"name": "E-Signature", "purpose": "Electronic signature integration"},
            "037": {"name": "Access Control", "purpose": "Agency-based access management"},
            "038": {"name": "User Management", "purpose": "Roles, profiles, and permissions"},
            "039": {"name": "Opportunity Security", "purpose": "Record-level opportunity access"},
            "040": {"name": "DMS Inventory Sync", "purpose": "Dealer Management System integration"},
            "041": {"name": "Invoice Exchange", "purpose": "Financial data synchronization"},
            "042": {"name": "Delivery Tracking", "purpose": "Delivery status management"},
            "043": {"name": "Conversion Metrics", "purpose": "Real-time lead conversion analytics"},
            "044": {"name": "Sales Alerts", "purpose": "Threshold-based sales notifications"},
            "045": {"name": "Pipeline Projections", "purpose": "Forecast trend analysis"},
            "046": {"name": "Language Preferences", "purpose": "Multi-language user settings"},
            "047": {"name": "Document Localization", "purpose": "Multi-language document generation"},
            "048": {"name": "Audit Trail", "purpose": "Change tracking and compliance"},
            "049": {"name": "Data Reversibility", "purpose": "Data export and restoration"},
            "050": {"name": "Training Management", "purpose": "User training and onboarding"},
            "051": {"name": "Sandbox Management", "purpose": "Environment and rollback procedures"},
        }
        
        # Group use cases by prefix
        uc_groups = defaultdict(list)
        for uc in use_cases:
            if isinstance(uc, dict):
                uc_id = uc.get("id", uc.get("use_case_id", ""))
                if uc_id and "-" in uc_id:
                    parts = uc_id.split("-")
                    if len(parts) >= 2:
                        prefix = parts[1]
                        uc_groups[prefix].append(uc)
        
        # Build modules list
        for prefix in sorted(uc_groups.keys()):
            module_info = module_defs.get(prefix, {"name": f"Module {prefix}", "purpose": "Functional module"})
            self.modules.append({
                "id": prefix,
                "name": module_info["name"],
                "purpose": module_info["purpose"],
                "use_cases": uc_groups[prefix]
            })
    
    def generate(self) -> str:
        """Generate the complete SDS document"""
        self.doc = Document()
        self._setup_styles()
        
        # Generate all sections
        self._add_title_page()
        self._add_document_control()
        self._add_table_of_contents()
        self._add_section_1_introduction()
        self._add_section_2_solution_scope()
        self._add_section_3_roles_profiles()
        self._add_section_4_functional_design()
        self._add_section_5_data_migration()
        self._add_section_6_interfaces()
        self._add_section_7_assumptions()
        self._add_section_8_signoff()
        
        # SDS-02/03/04: Add QA, Deployment, and Training sections
        self._add_section_9_quality_assurance()
        self._add_section_10_deployment()
        self._add_section_11_training()
        
        # Save document
        os.makedirs(OUTPUT_DIR, exist_ok=True)
        doc_id = f"SDS-{self.execution_id:04d}-{self.project_info['project_id']:02d}"
        filename = f"{doc_id}-{self.project_info['client_name'].replace(' ', '_')}.docx"
        output_path = os.path.join(OUTPUT_DIR, filename)
        self.doc.save(output_path)
        
        print(f"\n✅ SDS Document generated: {output_path}")
        return output_path
    
    def _setup_styles(self):
        """Setup document styles"""
        styles = self.doc.styles
        
        # Modify Heading 1
        h1 = styles['Heading 1']
        h1.font.color.rgb = HEADER_BLUE
        h1.font.size = Pt(16)
        
        # Modify Heading 2
        h2 = styles['Heading 2']
        h2.font.color.rgb = HEADER_BLUE
        h2.font.size = Pt(14)
        
        # Modify Heading 3
        h3 = styles['Heading 3']
        h3.font.color.rgb = LIGHT_BLUE
        h3.font.size = Pt(12)
    
    def _add_title_page(self):
        """Add title page with branding"""
        # Company header
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(COMPANY_NAME)
        run.font.size = Pt(11)
        run.font.color.rgb = HEADER_BLUE
        run.bold = True
        
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(COMPANY_ADDRESS)
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY_TEXT
        
        # Spacing
        for _ in range(4):
            self.doc.add_paragraph()
        
        # Main title
        doc_id = f"SDS-{self.execution_id:04d}-{self.project_info['project_id']:02d}-{self.project_info['client_name'].upper().replace(' ', '')}"
        
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Solution Design Specification")
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = HEADER_BLUE
        
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[{doc_id}]")
        run.font.size = Pt(14)
        run.font.color.rgb = GRAY_TEXT
        
        # Spacing
        for _ in range(3):
            self.doc.add_paragraph()
        
        # Metadata table
        table = self.doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        
        metadata = [
            ("Date", self.project_info["date"]),
            ("Client", self.project_info["client_name"]),
            ("Project", self.project_info["project_name"]),
            ("SOW Reference", f"SOW-{self.execution_id:04d}"),
            ("Author", "Digital Humans System")
        ]
        
        for i, (label, value) in enumerate(metadata):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            set_cell_shading(table.rows[i].cells[0], "E7E6E6")
            table.rows[i].cells[1].text = value
        
        self.doc.add_page_break()
    
    def _add_document_control(self):
        """Add document control section"""
        self.doc.add_heading("Document Control", 1)
        
        # Document Owner
        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        row = table.rows[0].cells
        row[0].text = "Document Owner"
        row[0].paragraphs[0].runs[0].bold = True
        set_cell_shading(row[0], "E7E6E6")
        row[1].text = "Digital Humans System"
        row[2].text = "contact@digital-humans.fr"
        row[3].text = ""
        
        self.doc.add_paragraph()
        
        # Document History
        p = self.doc.add_paragraph()
        run = p.add_run("Document History:")
        run.bold = True
        
        create_table_with_header(
            self.doc,
            ["Description", "Author", "Version", "Date"],
            [
                ["Initial Generation", "Digital Humans System", "1.0", self.project_info["date"]],
                ["", "", "", ""],
            ]
        )
        
        # Distribution List
        p = self.doc.add_paragraph()
        run = p.add_run("Distribution List:")
        run.bold = True
        
        create_table_with_header(
            self.doc,
            ["Name", "Organization", "Position"],
            [
                [self.project_info["client_name"], "", "Project Sponsor"],
                ["", "", ""],
            ]
        )
        
        self.doc.add_page_break()
    
    def _add_table_of_contents(self):
        """Add table of contents"""
        self.doc.add_heading("Contents", 1)
        
        toc_items = [
            ("1.", "Introduction", 1),
            ("2.", "Solution Scope", 1),
            ("", "2.1. Solution Overview", 2),
            ("", "2.2. Key Application Modules", 2),
            ("", "2.3. Key System Interfaces", 2),
            ("", "2.4. High Level Data Model", 2),
            ("3.", "System Roles & Profiles", 1),
            ("", "3.1. Key Roles", 2),
            ("", "3.2. Key Profiles", 2),
            ("4.", "Functional Design", 1),
        ]
        
        # Add modules to TOC
        for i, module in enumerate(self.modules[:10], 1):  # First 10 modules
            toc_items.append(("", f"4.{i}. Module: {module['name']}", 2))
        
        toc_items.extend([
            ("5.", "Data Migration Design", 1),
            ("6.", "Systems Interface Design", 1),
            ("7.", "Assumptions & Dependencies", 1),
            ("8.", "Acceptance and Signoff", 1),
        ])
        
        for num, title, level in toc_items:
            p = self.doc.add_paragraph()
            indent = "    " * (level - 1)
            run = p.add_run(f"{indent}{num} {title}" if num else f"{indent}{title}")
            if level == 1:
                run.bold = True
            run.font.size = Pt(11 if level == 1 else 10)
        
        self.doc.add_page_break()
    
    def _add_section_1_introduction(self):
        """Section 1: Introduction"""
        self.doc.add_heading("1. Introduction", 1)
        
        # Get project summary from PM (Sophie)
        pm_data = self.data.get("sophie", {})
        pm_content = pm_data.get("content", pm_data)
        project_summary = pm_content.get("project_summary", "")
        
        intro_text = f"""The purpose of this document is to outline the solution design for implementing {self.project_info['project_name']} for {self.project_info['client_name']}.

The contents of this document have been compiled following an analysis of functional requirements as defined in the Statement of Work and subsequent discussions and workshops held with client representatives."""
        
        self.doc.add_paragraph(intro_text)
        
        if project_summary:
            self.doc.add_paragraph()
            p = self.doc.add_paragraph()
            run = p.add_run("Project Overview: ")
            run.bold = True
            p.add_run(project_summary)
        
        self.doc.add_paragraph()
        p = self.doc.add_paragraph()
        run = p.add_run("This document will describe:")
        
        items = [
            "The overall scope of the solution including architecture overview",
            "Key User Roles and Profiles",
            "The functional design of each module including data model, workflows, and automation",
            "Data migration design and strategy",
            "Systems interface design for external integrations",
            "Design assumptions and dependencies"
        ]
        
        for item in items:
            p = self.doc.add_paragraph(f"• {item}")
            p.paragraph_format.left_indent = Inches(0.25)
        
        self.doc.add_page_break()
    
    def _add_section_2_solution_scope(self):
        """Section 2: Solution Scope"""
        self.doc.add_heading("2. Solution Scope", 1)
        
        marcus_data = self.data.get("marcus", {})
        content = marcus_data.get("content", {})
        
        # 2.1 Solution Overview
        self.doc.add_heading("2.1. Solution Overview", 2)
        
        title = content.get("title", self.project_info["project_name"])
        self.doc.add_paragraph(f"This solution implements {title} on the Salesforce {self.project_info['salesforce_product']} platform.")
        
        # Architecture description
        if content.get("technical_considerations"):
            p = self.doc.add_paragraph()
            run = p.add_run("Key Technical Components:")
            run.bold = True
            
            for tc in content.get("technical_considerations", [])[:5]:
                if isinstance(tc, dict):
                    category = tc.get("category", "")
                    consideration = tc.get("consideration", "")
                    p = self.doc.add_paragraph(f"• {category}: {consideration[:150]}...")
                    p.paragraph_format.left_indent = Inches(0.25)
        
        self.doc.add_paragraph()
        
        # 2.2 Key Application Modules
        self.doc.add_heading("2.2. Key Application Modules", 2)
        
        module_data = [[m["id"], m["name"], m["purpose"], str(len(m["use_cases"]))] for m in self.modules[:15]]
        create_table_with_header(
            self.doc,
            ["ID", "Module Name", "Purpose", "Use Cases"],
            module_data
        )
        
        # 2.3 Key System Interfaces
        self.doc.add_heading("2.3. Key System Interfaces", 2)
        
        integrations = content.get("integration_points", [])
        if integrations:
            int_data = []
            for integ in integrations[:10]:
                if isinstance(integ, dict):
                    int_data.append([
                        integ.get("system", ""),
                        integ.get("method", ""),
                        integ.get("direction", ""),
                        integ.get("frequency", "")[:30]
                    ])
            create_table_with_header(
                self.doc,
                ["System", "Method", "Direction", "Frequency"],
                int_data
            )
        else:
            self.doc.add_paragraph("System interfaces to be defined during detailed design phase.")
        
        # 2.4 High Level Data Model
        self.doc.add_heading("2.4. High Level Data Model", 2)
        
        data_model = content.get("data_model", {})
        erd_mermaid = data_model.get("erd_mermaid", "")
        
        if erd_mermaid:
            p = self.doc.add_paragraph()
            run = p.add_run("Entity Relationship Diagram (Mermaid notation):")
            run.bold = True
            
            # Add ERD code in monospace
            p = self.doc.add_paragraph()
            run = p.add_run(erd_mermaid)
            run.font.name = "Consolas"
            run.font.size = Pt(8)
        
        # Objects summary
        std_objects = data_model.get("standard_objects", [])
        custom_objects = data_model.get("custom_objects", [])
        
        if std_objects or custom_objects:
            self.doc.add_paragraph()
            p = self.doc.add_paragraph()
            run = p.add_run("Data Model Summary:")
            run.bold = True
            self.doc.add_paragraph(f"• Standard Objects: {len(std_objects)}")
            self.doc.add_paragraph(f"• Custom Objects: {len(custom_objects)}")
        
        self.doc.add_page_break()
    
    def _add_section_3_roles_profiles(self):
        """Section 3: System Roles & Profiles"""
        self.doc.add_heading("3. System Roles & Profiles", 1)
        
        marcus_data = self.data.get("marcus", {})
        content = marcus_data.get("content", {})
        security = content.get("security_model", content.get("security_design", {}))
        
        # 3.1 Key Roles
        self.doc.add_heading("3.1. Key Roles", 2)
        
        roles = security.get("roles", [])
        if roles:
            role_data = []
            for role in roles:
                if isinstance(role, dict):
                    role_data.append([
                        role.get("name", ""),
                        role.get("description", role.get("purpose", ""))[:80]
                    ])
            create_table_with_header(self.doc, ["Role Name", "Description"], role_data)
        else:
            p = self.doc.add_paragraph()
            run = p.add_run("Section à compléter")
            run.italic = True
            run.font.color.rgb = GRAY_TEXT
            self.doc.add_paragraph("Roles will be defined based on client organizational structure.")
        
        # 3.2 Key Profiles
        self.doc.add_heading("3.2. Key Profiles", 2)
        
        profiles = security.get("profiles", [])
        if profiles:
            prof_data = []
            for prof in profiles:
                if isinstance(prof, dict):
                    prof_data.append([
                        prof.get("name", ""),
                        prof.get("description", prof.get("purpose", ""))[:80]
                    ])
            create_table_with_header(self.doc, ["Profile Name", "Description"], prof_data)
        else:
            p = self.doc.add_paragraph()
            run = p.add_run("Section à compléter")
            run.italic = True
            run.font.color.rgb = GRAY_TEXT
            self.doc.add_paragraph("Profiles will be configured based on role requirements.")
        
        # Permission Sets
        perm_sets = security.get("permission_sets", [])
        if perm_sets:
            self.doc.add_heading("3.3. Permission Sets", 2)
            ps_data = []
            for ps in perm_sets:
                if isinstance(ps, dict):
                    ps_data.append([
                        ps.get("name", ""),
                        ps.get("purpose", ps.get("description", ""))[:80]
                    ])
            create_table_with_header(self.doc, ["Permission Set", "Purpose"], ps_data)
        
        self.doc.add_page_break()
    
    def _add_section_4_functional_design(self):
        """Section 4: Functional Design - Per Module"""
        self.doc.add_heading("4. Functional Design", 1)
        
        marcus_data = self.data.get("marcus", {})
        content = marcus_data.get("content", {})
        data_model = content.get("data_model", {})
        automation = content.get("automation_design", {})
        
        pm_data = self.data.get("sophie", {})
        pm_content = pm_data.get("content", pm_data)
        business_reqs = pm_content.get("business_requirements", [])
        
        # Process each module
        for idx, module in enumerate(self.modules[:12], 1):  # Limit to 12 modules
            self.doc.add_heading(f"4.{idx}. Module: {module['name']}", 2)
            
            # 4.x.1 Purpose
            self.doc.add_heading(f"4.{idx}.1. Purpose", 3)
            self.doc.add_paragraph(module["purpose"])
            
            # 4.x.2 Key Functional Requirements
            self.doc.add_heading(f"4.{idx}.2. Key Functional Requirements", 3)
            
            # Find related BRs
            module_brs = [br for br in business_reqs if isinstance(br, dict)][:3]
            if module_brs:
                br_data = []
                for br in module_brs:
                    br_data.append([
                        br.get("id", ""),
                        br.get("title", br.get("requirement", ""))[:60],
                        br.get("priority", "Medium")
                    ])
                create_table_with_header(self.doc, ["ID", "Requirement", "Priority"], br_data)
            else:
                self.doc.add_paragraph("Requirements defined in Statement of Work.")
            
            # 4.x.3 Key Users and Use Cases
            self.doc.add_heading(f"4.{idx}.3. Key Users and Use Cases", 3)
            
            uc_data = []
            for uc in module["use_cases"][:10]:
                if isinstance(uc, dict):
                    uc_data.append([
                        uc.get("id", uc.get("use_case_id", "")),
                        uc.get("title", uc.get("name", ""))[:50],
                        uc.get("actor", uc.get("primary_actor", ""))[:30]
                    ])
            
            if uc_data:
                create_table_with_header(self.doc, ["UC ID", "Use Case", "Primary Actor"], uc_data)
            
            if len(module["use_cases"]) > 10:
                self.doc.add_paragraph(f"... and {len(module['use_cases']) - 10} additional use cases.")
            
            # 4.x.4 Outline Business Processes
            self.doc.add_heading(f"4.{idx}.4. Outline Business Processes", 3)
            self.doc.add_paragraph(f"Business processes for {module['name']} are driven by the use cases defined above.")
            
            # 4.x.5 Data Model
            self.doc.add_heading(f"4.{idx}.5. Data Model", 3)
            
            # Find relevant objects
            custom_objects = data_model.get("custom_objects", [])
            if custom_objects and idx <= len(custom_objects):
                obj = custom_objects[min(idx-1, len(custom_objects)-1)]
                if isinstance(obj, dict):
                    p = self.doc.add_paragraph()
                    run = p.add_run(f"Primary Object: {obj.get('label', '')} ({obj.get('api_name', '')})")
                    run.bold = True
                    
                    fields = obj.get("fields", [])[:8]
                    if fields:
                        field_data = []
                        for f in fields:
                            if isinstance(f, str):
                                parts = f.split(" (", 1)
                                field_data.append([parts[0], parts[1].rstrip(")") if len(parts) > 1 else ""])
                        if field_data:
                            create_table_with_header(self.doc, ["Field", "Type"], field_data)
            else:
                self.doc.add_paragraph("Data model leverages standard Salesforce objects with custom fields.")
            
            # 4.x.6 Workflows
            self.doc.add_heading(f"4.{idx}.6. Workflows", 3)
            
            flows = automation.get("flows", [])
            if flows and idx <= len(flows):
                flow = flows[min(idx-1, len(flows)-1)]
                if isinstance(flow, dict):
                    self.doc.add_paragraph(f"• {flow.get('name', '')}: {flow.get('purpose', '')[:100]}")
            else:
                self.doc.add_paragraph("Automation to be defined during detailed design.")
            
            # 4.x.7 Apex, Triggers and Components
            self.doc.add_heading(f"4.{idx}.7. Apex, Triggers and Lightning Components", 3)
            
            triggers = automation.get("triggers", [])
            if triggers:
                for trig in triggers[:2]:
                    if isinstance(trig, dict):
                        self.doc.add_paragraph(f"• {trig.get('name', '')}: {trig.get('purpose', '')[:80]}")
            else:
                self.doc.add_paragraph("Custom code requirements to be identified during detailed design.")
            
            # 4.x.8 Reports & Dashboards
            self.doc.add_heading(f"4.{idx}.8. Reports & Dashboards", 3)
            self.doc.add_paragraph(f"Standard reports and dashboards for {module['name']} module.")
            
            # 4.x.9 3rd Party Applications
            self.doc.add_heading(f"4.{idx}.9. 3rd Party Applications", 3)
            self.doc.add_paragraph("No third-party applications required for this module.")
            
            # 4.x.10 Key System Interfaces
            self.doc.add_heading(f"4.{idx}.10. Key System Interfaces", 3)
            
            integrations = content.get("integration_points", [])
            if integrations and idx <= len(integrations):
                integ = integrations[min(idx-1, len(integrations)-1)]
                if isinstance(integ, dict):
                    self.doc.add_paragraph(f"• {integ.get('system', '')}: {integ.get('purpose', '')[:80]}")
            else:
                self.doc.add_paragraph("No specific interfaces required for this module.")
            
            # Page break between modules (except last)
            if idx < min(len(self.modules), 12):
                self.doc.add_page_break()
        
        self.doc.add_page_break()
    
    def _add_section_5_data_migration(self):
        """Section 5: Data Migration Design"""
        self.doc.add_heading("5. Data Migration Design", 1)
        
        # 5.1 Migration Strategy
        self.doc.add_heading("5.1. Migration Strategy & Plan", 2)
        p = self.doc.add_paragraph()
        run = p.add_run("Section à compléter")
        run.italic = True
        run.font.color.rgb = GRAY_TEXT
        self.doc.add_paragraph("Migration strategy will be defined by Aisha (Data Migration Specialist).")
        
        # 5.2 Key Data Sources
        self.doc.add_heading("5.2. Key Data Sources", 2)
        
        marcus_data = self.data.get("marcus", {})
        content = marcus_data.get("content", {})
        integrations = content.get("integration_points", [])
        
        # Extract potential data sources from integrations
        sources = []
        for integ in integrations:
            if isinstance(integ, dict):
                system = integ.get("system", "")
                if "DMS" in system or "Management" in system or "Legacy" in system.lower():
                    sources.append([system, integ.get("purpose", "")[:60], integ.get("method", "")])
        
        if sources:
            create_table_with_header(self.doc, ["Source System", "Data Type", "Method"], sources)
        else:
            p = self.doc.add_paragraph()
            run = p.add_run("Section à compléter")
            run.italic = True
            run.font.color.rgb = GRAY_TEXT
            self.doc.add_paragraph("Data sources will be identified during migration analysis phase.")
        
        # 5.3 Data Scope
        self.doc.add_heading("5.3. Data Scope", 2)
        
        data_model = content.get("data_model", {})
        custom_objects = data_model.get("custom_objects", [])
        
        if custom_objects:
            scope_data = []
            for obj in custom_objects[:6]:
                if isinstance(obj, dict):
                    scope_data.append([
                        obj.get("api_name", ""),
                        obj.get("label", ""),
                        "To be defined"
                    ])
            if scope_data:
                create_table_with_header(self.doc, ["Object", "Description", "Scope"], scope_data)
            else:
                p = self.doc.add_paragraph()
                run = p.add_run("Section à compléter")
                run.italic = True
                run.font.color.rgb = GRAY_TEXT
        else:
            p = self.doc.add_paragraph()
            run = p.add_run("Section à compléter")
            run.italic = True
            run.font.color.rgb = GRAY_TEXT
            self.doc.add_paragraph("Data scope will be defined during migration analysis phase.")
        
        # 5.4 Migration Designs placeholder
        self.doc.add_heading("5.4. Migration Designs", 2)
        
        p = self.doc.add_paragraph()
        run = p.add_run("Note: ")
        run.bold = True
        run.font.color.rgb = LIGHT_BLUE
        p.add_run("Detailed migration file specifications will be developed by Aisha (Data Migration Specialist) during the detailed design phase. Each migration file will include:")
        
        items = [
            "Data Objects and field mappings",
            "Transformation and validation rules",
            "Importation method and tools",
            "Control checks and verification procedures",
            "Unmapped/bad data handling procedures"
        ]
        for item in items:
            p = self.doc.add_paragraph(f"• {item}")
            p.paragraph_format.left_indent = Inches(0.25)
        
        self.doc.add_page_break()
    
    def _add_section_6_interfaces(self):
        """Section 6: Systems Interface Design"""
        self.doc.add_heading("6. Systems Interface Design", 1)
        
        marcus_data = self.data.get("marcus", {})
        content = marcus_data.get("content", {})
        integrations = content.get("integration_points", [])
        
        # 6.1 Core Systems Interfaces
        self.doc.add_heading("6.1. Core Systems Interfaces", 2)
        
        if integrations:
            int_summary = []
            for integ in integrations:
                if isinstance(integ, dict):
                    int_summary.append([
                        integ.get("system", ""),
                        integ.get("direction", ""),
                        integ.get("method", ""),
                        integ.get("frequency", "")[:25]
                    ])
            create_table_with_header(
                self.doc,
                ["System", "Direction", "Method", "Frequency"],
                int_summary
            )
            
            # Detailed interface sections
            for idx, integ in enumerate(integrations[:6], 1):
                if isinstance(integ, dict):
                    system_name = integ.get("system", f"Interface {idx}")
                    self.doc.add_heading(f"6.1.{idx}. {system_name}", 3)
                    
                    # Create detail table
                    table = self.doc.add_table(rows=6, cols=2)
                    table.style = 'Table Grid'
                    
                    details = [
                        ("Interfacing Application", integ.get("system", "")),
                        ("Business Processes", integ.get("purpose", "")),
                        ("Integration Method", integ.get("method", "")),
                        ("Data Mapping & Rules", integ.get("authentication", "") + " authentication"),
                        ("Timing", integ.get("frequency", "")),
                        ("Exception Handling", integ.get("error_handling", ""))
                    ]
                    
                    for i, (label, value) in enumerate(details):
                        table.rows[i].cells[0].text = label
                        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
                        set_cell_shading(table.rows[i].cells[0], "E7E6E6")
                        table.rows[i].cells[1].text = value[:150] if value else "To be defined"
                    
                    self.doc.add_paragraph()
        else:
            self.doc.add_paragraph("System interfaces to be defined during detailed design phase.")
        
        self.doc.add_page_break()
    
    def _add_section_7_assumptions(self):
        """Section 7: Assumptions & Dependencies"""
        self.doc.add_heading("7. Assumptions & Dependencies", 1)
        
        marcus_data = self.data.get("marcus", {})
        content = marcus_data.get("content", {})
        
        # 7.1 Key Design Assumptions
        self.doc.add_heading("7.1. Key Design Assumptions", 2)
        
        # Try to get assumptions from Marcus data
        assumptions = content.get("assumptions", [])
        
        if assumptions:
            for assumption in assumptions:
                if isinstance(assumption, dict):
                    self.doc.add_paragraph(f"• {assumption.get('assumption', assumption.get('description', ''))}")
                elif isinstance(assumption, str):
                    self.doc.add_paragraph(f"• {assumption}")
        else:
            p = self.doc.add_paragraph()
            run = p.add_run("Section à compléter")
            run.italic = True
            run.font.color.rgb = GRAY_TEXT
            self.doc.add_paragraph("Design assumptions will be documented during detailed design phase.")
        
        # 7.2 Key Dependencies
        self.doc.add_heading("7.2. Key Dependencies", 2)
        
        # Extract from technical considerations
        tech_considerations = content.get("technical_considerations", [])
        risks = content.get("risks", [])
        
        dependencies = []
        
        for tc in tech_considerations[:5]:
            if isinstance(tc, dict):
                dependencies.append(f"{tc.get('category', '')}: {tc.get('consideration', '')[:100]}")
        
        if dependencies:
            for dep in dependencies:
                self.doc.add_paragraph(f"• {dep}")
        else:
            p = self.doc.add_paragraph()
            run = p.add_run("Section à compléter")
            run.italic = True
            run.font.color.rgb = GRAY_TEXT
            self.doc.add_paragraph("Dependencies will be identified during project planning phase.")
        
        # Add risks section
        if risks:
            self.doc.add_heading("7.3. Key Risks", 2)
            
            risk_data = []
            for risk in risks[:8]:
                if isinstance(risk, dict):
                    risk_data.append([
                        risk.get("risk", ""),
                        risk.get("severity", ""),
                        risk.get("mitigation", "")[:60]
                    ])
            
            if risk_data:
                create_table_with_header(self.doc, ["Risk", "Severity", "Mitigation"], risk_data)
        
        self.doc.add_page_break()
    
    def _add_section_8_signoff(self):
        """Section 8: Acceptance and Signoff"""
        self.doc.add_heading("8. Acceptance and Signoff", 1)
        
        self.doc.add_paragraph("We accept this document to form the basis of our solution design.")
        
        self.doc.add_paragraph()
        
        # Signoff table
        table = self.doc.add_table(rows=6, cols=3)
        table.style = 'Table Grid'
        
        table.rows[0].cells[0].text = self.project_info["client_name"]
        table.rows[0].cells[0].merge(table.rows[0].cells[2])
        table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[0], "E7E6E6")
        
        labels = ["By:", "", "Name:", "Title:", "Date:"]
        for i, label in enumerate(labels, 1):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        # Footer
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Generated by {COMPANY_NAME}")
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY_TEXT
    # =========================================================================
    # SDS-02: SECTION 9 - QUALITY ASSURANCE (Elena)
    # =========================================================================
    def _add_section_9_quality_assurance(self):
        """Section 9: Quality Assurance - Test Strategy and Test Cases"""
        self.doc.add_heading("9. Quality Assurance", 1)
        
        # Get QA data from Elena's artifacts
        qa_data = self.data.get("elena", {})
        content = qa_data.get("content", qa_data)
        
        # 9.1 Test Strategy
        self.doc.add_heading("9.1. Test Strategy", 2)
        test_strategy = content.get("test_strategy", {})
        
        if test_strategy:
            approach = test_strategy.get("approach", "Comprehensive testing covering unit, integration, and UAT")
            self.doc.add_paragraph(approach)
            
            # Test levels
            self.doc.add_heading("Test Levels", 3)
            levels = test_strategy.get("test_levels", [
                {"level": "Unit Testing", "responsibility": "Developers", "coverage": "85%+"},
                {"level": "Integration Testing", "responsibility": "QA Team", "coverage": "All integrations"},
                {"level": "UAT", "responsibility": "Business Users", "coverage": "Critical paths"}
            ])
            if levels:
                data = [[l.get("level", ""), l.get("responsibility", ""), l.get("coverage", "")] for l in levels]
                create_table_with_header(self.doc, ["Test Level", "Responsibility", "Coverage Target"], data)
        else:
            self.doc.add_paragraph("Test strategy to be defined during project planning.")
        
        # 9.2 Test Cases Summary
        self.doc.add_heading("9.2. Test Cases Summary", 2)
        test_cases = content.get("test_cases", content.get("tests", []))
        
        if test_cases:
            self.doc.add_paragraph(f"Total test cases: {len(test_cases)}")
            
            # Group by type
            by_type = {}
            for tc in test_cases:
                tc_type = tc.get("type", tc.get("category", "General"))
                if tc_type not in by_type:
                    by_type[tc_type] = []
                by_type[tc_type].append(tc)
            
            # Summary table
            summary_data = [[t, str(len(cases))] for t, cases in by_type.items()]
            if summary_data:
                create_table_with_header(self.doc, ["Test Type", "Count"], summary_data)
            
            # Detailed test cases (max 50 to avoid huge document)
            self.doc.add_heading("9.3. Detailed Test Cases", 2)
            
            test_data = []
            for tc in test_cases[:50]:
                test_data.append([
                    tc.get("id", tc.get("test_id", "")),
                    tc.get("name", tc.get("title", ""))[:40],
                    tc.get("type", tc.get("category", "")),
                    tc.get("uc_ref", tc.get("related_uc", ""))
                ])
            
            if test_data:
                create_table_with_header(self.doc, ["Test ID", "Name", "Type", "UC Reference"], test_data)
                
                if len(test_cases) > 50:
                    p = self.doc.add_paragraph()
                    run = p.add_run(f"... and {len(test_cases) - 50} more test cases (see appendix)")
                    run.italic = True
                    run.font.color.rgb = GRAY_TEXT
        else:
            p = self.doc.add_paragraph()
            run = p.add_run("Test cases will be developed during the build phase.")
            run.italic = True
            run.font.color.rgb = GRAY_TEXT
        
        self.doc.add_page_break()
    
    # =========================================================================
    # SDS-03: SECTION 10 - DEPLOYMENT (Jordan)
    # =========================================================================
    def _add_section_10_deployment(self):
        """Section 10: Deployment Plan - CI/CD and Release Strategy"""
        self.doc.add_heading("10. Deployment Plan", 1)
        
        # Get DevOps data from Jordan's artifacts
        devops_data = self.data.get("jordan", {})
        content = devops_data.get("content", devops_data)
        
        # 10.1 Deployment Strategy
        self.doc.add_heading("10.1. Deployment Strategy", 2)
        strategy = content.get("deployment_strategy", content.get("strategy", {}))
        
        approach = strategy.get("approach", "Metadata API deployment with scratch org validation")
        self.doc.add_paragraph(approach)
        
        # Environments
        self.doc.add_heading("Environments", 3)
        environments = content.get("environments", [
            {"name": "Developer Sandbox", "type": "Developer", "purpose": "Individual development"},
            {"name": "Integration Sandbox", "type": "Developer Pro", "purpose": "Integration testing"},
            {"name": "UAT Sandbox", "type": "Partial Copy", "purpose": "User acceptance testing"},
            {"name": "Production", "type": "Production", "purpose": "Live environment"}
        ])
        
        if environments:
            env_data = [[e.get("name", ""), e.get("type", ""), e.get("purpose", "")] for e in environments]
            create_table_with_header(self.doc, ["Environment", "Type", "Purpose"], env_data)
        
        # 10.2 CI/CD Pipeline
        self.doc.add_heading("10.2. CI/CD Pipeline", 2)
        pipeline = content.get("cicd_pipeline", content.get("pipeline", {}))
        
        stages = pipeline.get("stages", [
            {"stage": "Build", "actions": "Source validation, static analysis"},
            {"stage": "Test", "actions": "Apex tests, Jest tests"},
            {"stage": "Deploy", "actions": "Deploy to target org"},
            {"stage": "Validate", "actions": "Smoke tests, data validation"}
        ])
        
        if stages:
            stage_data = [[s.get("stage", ""), s.get("actions", "")] for s in stages]
            create_table_with_header(self.doc, ["Stage", "Actions"], stage_data)
        
        # 10.3 Pre-Deployment Checklist
        self.doc.add_heading("10.3. Pre-Deployment Checklist", 2)
        checklist = content.get("pre_deployment_checklist", content.get("checklist", []))
        
        if checklist:
            for item in checklist[:15]:
                if isinstance(item, dict):
                    self.doc.add_paragraph(f"□ {item.get('item', item.get('description', ''))}")
                else:
                    self.doc.add_paragraph(f"□ {item}")
        else:
            default_checklist = [
                "All unit tests pass (85%+ coverage)",
                "Code review completed and approved",
                "UAT sign-off obtained",
                "Release notes prepared",
                "Rollback plan documented",
                "Communication sent to stakeholders"
            ]
            for item in default_checklist:
                self.doc.add_paragraph(f"□ {item}")
        
        # 10.4 Rollback Plan
        self.doc.add_heading("10.4. Rollback Plan", 2)
        rollback = content.get("rollback_plan", {})
        
        if rollback:
            self.doc.add_paragraph(rollback.get("description", ""))
            steps = rollback.get("steps", [])
            for i, step in enumerate(steps, 1):
                self.doc.add_paragraph(f"{i}. {step}")
        else:
            self.doc.add_paragraph("In case of deployment issues, the following rollback procedure will be executed:")
            default_steps = [
                "Stop ongoing deployment immediately",
                "Notify all stakeholders of rollback initiation",
                "Deploy previous stable package version",
                "Verify system functionality with smoke tests",
                "Document issues for root cause analysis"
            ]
            for i, step in enumerate(default_steps, 1):
                self.doc.add_paragraph(f"{i}. {step}")
        
        self.doc.add_page_break()
    
    # =========================================================================
    # SDS-04: SECTION 11 - TRAINING & ADOPTION (Lucas)
    # =========================================================================
    def _add_section_11_training(self):
        """Section 11: Training & Adoption Strategy"""
        self.doc.add_heading("11. Training & Adoption", 1)
        
        # Get Trainer data from Lucas's artifacts
        trainer_data = self.data.get("lucas", {})
        content = trainer_data.get("content", trainer_data)
        
        # 11.1 Training Strategy
        self.doc.add_heading("11.1. Training Strategy", 2)
        strategy = content.get("training_strategy", content.get("strategy", {}))
        
        summary = strategy.get("executive_summary", content.get("executive_summary", ""))
        if summary:
            self.doc.add_paragraph(summary)
        
        # Training approach
        approach = content.get("training_approach", strategy.get("approach", {}))
        if approach:
            methodology = approach.get("methodology", "Blended learning approach")
            self.doc.add_paragraph(f"Methodology: {methodology}")
            
            methods = approach.get("delivery_methods", [])
            if methods:
                self.doc.add_paragraph("Delivery Methods:")
                for method in methods:
                    self.doc.add_paragraph(f"• {method}")
        
        # 11.2 Audience Analysis
        self.doc.add_heading("11.2. Audience Analysis", 2)
        audience = content.get("audience_analysis", {})
        personas = audience.get("user_personas", content.get("user_personas", []))
        
        if personas:
            persona_data = []
            for p in personas[:10]:
                persona_data.append([
                    p.get("role", ""),
                    p.get("count_estimate", ""),
                    p.get("sf_experience", ""),
                    p.get("training_priority", "")
                ])
            create_table_with_header(self.doc, ["Role", "Est. Users", "SF Experience", "Priority"], persona_data)
        else:
            self.doc.add_paragraph("User personas to be identified during discovery phase.")
        
        # 11.3 Training Curriculum
        self.doc.add_heading("11.3. Training Curriculum", 2)
        curriculum = content.get("curriculum_outline", content.get("curriculum", []))
        
        if curriculum:
            curr_data = []
            for c in curriculum[:15]:
                curr_data.append([
                    c.get("module", ""),
                    c.get("duration", ""),
                    c.get("audience", "")
                ])
            create_table_with_header(self.doc, ["Module", "Duration", "Audience"], curr_data)
        else:
            self.doc.add_paragraph("Training modules to be defined based on user analysis.")
        
        # 11.4 Adoption Metrics
        self.doc.add_heading("11.4. Adoption Metrics", 2)
        metrics = content.get("adoption_metrics", {})
        kpis = metrics.get("kpis", content.get("kpis", []))
        
        if kpis:
            kpi_data = []
            for k in kpis[:10]:
                kpi_data.append([
                    k.get("metric", ""),
                    k.get("target", ""),
                    k.get("measurement", "")
                ])
            create_table_with_header(self.doc, ["Metric", "Target", "Measurement"], kpi_data)
        else:
            default_kpis = [
                {"metric": "Login Rate", "target": ">90% weekly", "measurement": "Salesforce Reports"},
                {"metric": "Data Quality", "target": ">95% complete", "measurement": "Validation Rules"},
                {"metric": "Feature Adoption", "target": ">80% in 30 days", "measurement": "Usage Analytics"}
            ]
            kpi_data = [[k["metric"], k["target"], k["measurement"]] for k in default_kpis]
            create_table_with_header(self.doc, ["Metric", "Target", "Measurement"], kpi_data)
        
        # 11.5 Timeline
        self.doc.add_heading("11.5. Training Timeline", 2)
        timeline = content.get("timeline", {})
        
        if timeline:
            phases = [
                f"Preparation: {timeline.get('prep_weeks', 2)} weeks",
                f"Delivery: {timeline.get('delivery_weeks', 4)} weeks",
                f"Reinforcement: {timeline.get('reinforcement_weeks', 4)} weeks"
            ]
            for phase in phases:
                self.doc.add_paragraph(f"• {phase}")
            
            milestones = timeline.get("key_milestones", [])
            if milestones:
                self.doc.add_heading("Key Milestones", 3)
                for m in milestones:
                    self.doc.add_paragraph(f"• {m}")
        else:
            self.doc.add_paragraph("Training timeline to be established during project planning.")
        
        self.doc.add_page_break()



def generate_sds_from_template(execution_id: int) -> str:
    """Main function to generate SDS document"""
    generator = SDSTemplateGenerator(execution_id)
    generator.load_data()
    return generator.generate()


if __name__ == "__main__":
    import sys
    exec_id = int(sys.argv[1]) if len(sys.argv) > 1 else 79
    output = generate_sds_from_template(exec_id)
    print(f"\n✅ Document generated: {output}")
