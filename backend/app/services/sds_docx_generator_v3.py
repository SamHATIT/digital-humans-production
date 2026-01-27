"""
SDS DOCX Generator v3 - Génération de documents Word professionnels

Convertit le résultat de SDSSynthesisService en document Word de qualité consulting
Style inspiré des SDS LVMH (86p) et Shiseido/Itelios (73p)

Caractéristiques:
- Page de titre professionnelle
- Table des matières automatique
- Numérotation hiérarchique des sections
- Tableaux formatés avec en-têtes colorés
- Diagrammes Mermaid convertis en images
- En-têtes/pieds de page avec pagination
"""

import os
import re
import io
import base64
import logging
import tempfile
import subprocess
from typing import Dict, List, Any, Optional, Tuple
from datetime import datetime
from dataclasses import dataclass

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

logger = logging.getLogger(__name__)


# ============================================================================
# CONSTANTES DE STYLE (inspiré Itelios/Shiseido)
# ============================================================================

# Couleurs corporate
PRIMARY_BLUE = RGBColor(31, 78, 121)      # #1F4E79 - Titres principaux
SECONDARY_BLUE = RGBColor(68, 114, 196)   # #4472C4 - Sous-titres
ACCENT_ORANGE = RGBColor(237, 125, 49)    # #ED7D31 - Accents (comme Itelios)
TABLE_HEADER_BG = "1F4E79"                 # Fond en-têtes tableaux
TABLE_ALT_ROW = "F2F2F2"                   # Lignes alternées
LINK_BLUE = RGBColor(37, 99, 235)         # #2563EB
GRAY_TEXT = RGBColor(107, 114, 128)       # Texte secondaire

# Polices
FONT_TITLE = "Calibri Light"
FONT_BODY = "Calibri"
FONT_CODE = "Consolas"

# Tailles
SIZE_TITLE = Pt(28)
SIZE_H1 = Pt(16)
SIZE_H2 = Pt(14)
SIZE_H3 = Pt(12)
SIZE_BODY = Pt(11)
SIZE_SMALL = Pt(9)
SIZE_CODE = Pt(9)


# ============================================================================
# CLASSE PRINCIPALE
# ============================================================================

class SDSDocxGeneratorV3:
    """Génère des documents SDS Word de qualité professionnelle"""
    
    def __init__(self):
        self.doc = None
        self.section_counter = [0, 0, 0, 0]  # Pour numérotation hiérarchique
        self.toc_entries = []  # Pour table des matières manuelle si besoin
        
    def generate(
        self,
        project_name: str,
        synthesis_result: Dict[str, Any],
        wbs_data: Optional[Dict[str, Any]] = None,
        project_info: Optional[Dict[str, Any]] = None,
        output_path: Optional[str] = None
    ) -> bytes:
        """
        Génère le document SDS complet
        
        Args:
            project_name: Nom du projet
            synthesis_result: Résultat de SDSSynthesisService.synthesize_sds()
            wbs_data: Données WBS de Marcus (optionnel)
            project_info: Infos projet additionnelles (product, org_type, etc.)
            output_path: Chemin de sortie (optionnel, sinon retourne bytes)
            
        Returns:
            bytes du document DOCX ou chemin du fichier
        """
        self.doc = Document()
        self._setup_document()
        self._setup_styles()
        
        project_info = project_info or {}
        
        # 1. Page de titre
        self._add_title_page(
            project_name=project_name,
            product=project_info.get("salesforce_product", "Salesforce Service Cloud"),
            org_type=project_info.get("organization_type", "Enterprise"),
            version="1.0",
            author="Digital Humans"
        )
        
        # 2. Table des matières
        self._add_table_of_contents()
        
        # 3. Executive Summary
        self._add_executive_summary(
            project_name=project_name,
            total_ucs=synthesis_result.get("total_ucs", 0),
            domains_count=synthesis_result.get("domains_count", 0),
            sections=synthesis_result.get("sections", []),
            stats=synthesis_result.get("stats", {})
        )
        
        # 4. Sections par domaine fonctionnel
        for section in synthesis_result.get("sections", []):
            self._add_domain_section(section)
        
        # 5. Modèle de données global (ERD)
        if synthesis_result.get("erd_mermaid"):
            self._add_erd_section(synthesis_result["erd_mermaid"])
        
        # 6. Matrice des permissions
        if synthesis_result.get("permissions_matrix"):
            self._add_permissions_matrix(synthesis_result["permissions_matrix"])
        
        # 7. WBS / Plan de projet (si disponible)
        if wbs_data:
            self._add_wbs_section(wbs_data)
        
        # 8. Annexes (métadonnées de génération)
        self._add_appendix(synthesis_result)
        
        # Sauvegarde
        if output_path:
            self.doc.save(output_path)
            return output_path
        else:
            buffer = io.BytesIO()
            self.doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
    
    # ========================================================================
    # CONFIGURATION DOCUMENT
    # ========================================================================
    
    def _setup_document(self):
        """Configure les propriétés du document"""
        # Marges (style Itelios: marges étroites pour plus de contenu)
        for section in self.doc.sections:
            section.page_width = Inches(8.27)   # A4
            section.page_height = Inches(11.69)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2)
            
            # En-tête et pied de page
            section.header_distance = Cm(1.25)
            section.footer_distance = Cm(1.25)
    
    def _setup_styles(self):
        """Configure les styles personnalisés"""
        styles = self.doc.styles
        
        # Style Heading 1
        h1_style = styles['Heading 1']
        h1_style.font.name = FONT_TITLE
        h1_style.font.size = SIZE_H1
        h1_style.font.bold = True
        h1_style.font.color.rgb = PRIMARY_BLUE
        h1_style.paragraph_format.space_before = Pt(18)
        h1_style.paragraph_format.space_after = Pt(6)
        h1_style.paragraph_format.keep_with_next = True
        
        # Style Heading 2
        h2_style = styles['Heading 2']
        h2_style.font.name = FONT_TITLE
        h2_style.font.size = SIZE_H2
        h2_style.font.bold = True
        h2_style.font.color.rgb = SECONDARY_BLUE
        h2_style.paragraph_format.space_before = Pt(12)
        h2_style.paragraph_format.space_after = Pt(6)
        
        # Style Heading 3
        h3_style = styles['Heading 3']
        h3_style.font.name = FONT_BODY
        h3_style.font.size = SIZE_H3
        h3_style.font.bold = True
        h3_style.font.color.rgb = ACCENT_ORANGE
        h3_style.paragraph_format.space_before = Pt(10)
        h3_style.paragraph_format.space_after = Pt(4)
        
        # Style Normal
        normal_style = styles['Normal']
        normal_style.font.name = FONT_BODY
        normal_style.font.size = SIZE_BODY
        normal_style.paragraph_format.space_after = Pt(6)
        normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        
        # Style Code (créer si n'existe pas)
        if 'Code' not in [s.name for s in styles]:
            code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
            code_style.font.name = FONT_CODE
            code_style.font.size = SIZE_CODE
            code_style.paragraph_format.space_before = Pt(6)
            code_style.paragraph_format.space_after = Pt(6)
            code_style.paragraph_format.left_indent = Cm(0.5)
    
    # ========================================================================
    # SECTIONS DU DOCUMENT
    # ========================================================================
    
    def _add_title_page(
        self, 
        project_name: str, 
        product: str, 
        org_type: str,
        version: str,
        author: str
    ):
        """Ajoute la page de titre style consulting"""
        
        # Espace en haut
        for _ in range(3):
            self.doc.add_paragraph()
        
        # Titre principal
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("Solution Design Specification")
        run.font.name = FONT_TITLE
        run.font.size = SIZE_TITLE
        run.font.color.rgb = PRIMARY_BLUE
        run.font.bold = True
        
        # Ligne décorative
        line = self.doc.add_paragraph()
        line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = line.add_run("─" * 40)
        run.font.color.rgb = ACCENT_ORANGE
        
        # Nom du projet
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(project_name)
        run.font.name = FONT_TITLE
        run.font.size = Pt(22)
        run.font.bold = True
        
        # Produit Salesforce
        for _ in range(2):
            self.doc.add_paragraph()
        
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{product}")
        run.font.size = Pt(14)
        run.font.color.rgb = SECONDARY_BLUE
        
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{org_type}")
        run.font.size = Pt(12)
        run.font.color.rgb = GRAY_TEXT
        
        # Informations en bas
        for _ in range(8):
            self.doc.add_paragraph()
        
        # Tableau d'informations
        info_table = self.doc.add_table(rows=4, cols=2)
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        info_data = [
            ("Version", version),
            ("Date", datetime.now().strftime("%d/%m/%Y")),
            ("Auteur", author),
            ("Statut", "Draft")
        ]
        
        for i, (label, value) in enumerate(info_data):
            row = info_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            # Style
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.size = Pt(10)
        
        # Saut de page
        self.doc.add_page_break()
    
    def _add_table_of_contents(self):
        """Ajoute une table des matières"""
        self.doc.add_heading("Table des matières", 1)
        
        # Instruction pour mise à jour manuelle
        p = self.doc.add_paragraph()
        p.add_run("(Clic droit → Mettre à jour le champ pour actualiser)").font.italic = True
        p.runs[0].font.size = SIZE_SMALL
        p.runs[0].font.color.rgb = GRAY_TEXT
        
        # Champ TOC Word
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()
        fld_char_begin = parse_xml(r'<w:fldChar w:fldCharType="begin" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        instr_text = parse_xml(r'<w:instrText xml:space="preserve" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"> TOC \o "1-3" \h \z \u </w:instrText>')
        fld_char_separate = parse_xml(r'<w:fldChar w:fldCharType="separate" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        fld_char_end = parse_xml(r'<w:fldChar w:fldCharType="end" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        
        run._r.append(fld_char_begin)
        run._r.append(instr_text)
        run._r.append(fld_char_separate)
        run._r.append(fld_char_end)
        
        self.doc.add_page_break()
    
    def _add_executive_summary(
        self,
        project_name: str,
        total_ucs: int,
        domains_count: int,
        sections: List[Dict],
        stats: Dict
    ):
        """Ajoute le résumé exécutif"""
        self.doc.add_heading("1. Executive Summary", 1)
        
        # Paragraphe d'introduction
        intro = f"""Ce document présente la spécification technique détaillée (SDS) pour le projet {project_name}. 
Il décrit l'architecture de la solution, le modèle de données, les règles métier, et les automatisations 
nécessaires à l'implémentation sur la plateforme Salesforce."""
        
        self.doc.add_paragraph(intro)
        
        # Statistiques clés
        self.doc.add_heading("1.1 Périmètre du projet", 2)
        
        stats_table = self.doc.add_table(rows=4, cols=2)
        stats_table.style = 'Table Grid'
        self._style_table(stats_table, has_header=False)
        
        stats_data = [
            ("Domaines fonctionnels", str(domains_count)),
            ("Use Cases analysés", str(total_ucs)),
            ("Coût de génération", f"${stats.get('total_cost_usd', 0):.2f}"),
            ("Temps de génération", f"{stats.get('generation_time_ms', 0) / 1000:.1f}s")
        ]
        
        for i, (label, value) in enumerate(stats_data):
            row = stats_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            row.cells[0].paragraphs[0].runs[0].font.bold = True
        
        # Liste des domaines
        if sections:
            self.doc.add_heading("1.2 Domaines couverts", 2)
            for section in sections:
                p = self.doc.add_paragraph(style='List Bullet')
                run = p.add_run(f"{section.get('domain', 'N/A')}")
                run.font.bold = True
                p.add_run(f" ({section.get('uc_count', 0)} Use Cases)")
        
        self.doc.add_page_break()
    
    def _add_domain_section(self, section: Dict[str, Any]):
        """Ajoute une section de domaine fonctionnel"""
        domain = section.get("domain", "Unknown")
        content = section.get("content", "")
        
        # Convertir le Markdown en éléments Word
        self._parse_markdown_to_docx(content)
        
        self.doc.add_page_break()
    
    def _add_erd_section(self, erd_mermaid: str):
        """Ajoute la section ERD"""
        self.doc.add_heading("Modèle de données (ERD)", 1)
        
        # Essayer de convertir le Mermaid en image
        image_bytes = self._mermaid_to_image(erd_mermaid)
        
        if image_bytes:
            # Ajouter l'image
            image_stream = io.BytesIO(image_bytes)
            self.doc.add_picture(image_stream, width=Inches(6))
            last_paragraph = self.doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            # Fallback: afficher le code
            self.doc.add_paragraph("Diagramme ERD (format Mermaid):")
            p = self.doc.add_paragraph()
            p.style = 'Code' if 'Code' in [s.name for s in self.doc.styles] else 'Normal'
            p.add_run(erd_mermaid[:2000])  # Limiter la taille
        
        self.doc.add_page_break()
    
    def _add_permissions_matrix(self, matrix_content: str):
        """Ajoute la matrice des permissions"""
        self.doc.add_heading("Matrice des permissions", 1)
        
        # Parser le tableau Markdown
        lines = matrix_content.strip().split("\n")
        if len(lines) >= 2:
            # Extraire headers et rows
            headers = [h.strip() for h in lines[0].split("|") if h.strip()]
            rows = []
            for line in lines[2:]:  # Skip la ligne de séparation
                cells = [c.strip() for c in line.split("|") if c.strip()]
                if cells:
                    rows.append(cells)
            
            if headers and rows:
                table = self.doc.add_table(rows=len(rows)+1, cols=len(headers))
                table.style = 'Table Grid'
                self._style_table(table, has_header=True)
                
                # Headers
                for j, header in enumerate(headers):
                    table.rows[0].cells[j].text = header
                
                # Data
                for i, row_data in enumerate(rows):
                    for j, cell_data in enumerate(row_data):
                        if j < len(table.rows[i+1].cells):
                            table.rows[i+1].cells[j].text = cell_data
        else:
            self.doc.add_paragraph(matrix_content)
        
        self.doc.add_page_break()
    
    def _add_wbs_section(self, wbs_data: Dict[str, Any]):
        """Ajoute la section WBS / Plan de projet"""
        self.doc.add_heading("Plan de projet (WBS)", 1)
        
        phases = wbs_data.get("phases", [])
        if not phases and isinstance(wbs_data, dict):
            # Essayer d'extraire les phases autrement
            phases = wbs_data.get("content", {}).get("phases", [])
        
        if phases:
            for phase in phases:
                phase_name = phase.get("name", phase.get("phase", "Phase"))
                self.doc.add_heading(phase_name, 2)
                
                tasks = phase.get("tasks", phase.get("activities", []))
                if tasks:
                    # Créer un tableau pour les tâches
                    table = self.doc.add_table(rows=len(tasks)+1, cols=4)
                    table.style = 'Table Grid'
                    self._style_table(table, has_header=True)
                    
                    headers = ["Tâche", "Assigné", "Durée", "Dépendances"]
                    for j, h in enumerate(headers):
                        table.rows[0].cells[j].text = h
                    
                    for i, task in enumerate(tasks):
                        row = table.rows[i+1]
                        row.cells[0].text = task.get("name", task.get("task", ""))[:50]
                        row.cells[1].text = task.get("assigned_to", task.get("agent", ""))
                        row.cells[2].text = task.get("duration", task.get("effort", ""))
                        deps = task.get("dependencies", [])
                        row.cells[3].text = ", ".join(deps) if isinstance(deps, list) else str(deps)
        else:
            self.doc.add_paragraph("WBS non disponible ou format non reconnu.")
    
    def _add_appendix(self, synthesis_result: Dict):
        """Ajoute les annexes avec métadonnées"""
        self.doc.add_heading("Annexes", 1)
        
        self.doc.add_heading("A. Informations de génération", 2)
        
        info_table = self.doc.add_table(rows=5, cols=2)
        info_table.style = 'Table Grid'
        
        info_data = [
            ("Date de génération", synthesis_result.get("generated_at", datetime.now().isoformat())),
            ("Tokens utilisés", str(synthesis_result.get("stats", {}).get("total_tokens", 0))),
            ("Coût total", f"${synthesis_result.get('stats', {}).get('total_cost_usd', 0):.4f}"),
            ("Temps de génération", f"{synthesis_result.get('stats', {}).get('generation_time_ms', 0)}ms"),
            ("Version SDS", "v3 (Micro-analyse + Synthèse)")
        ]
        
        for i, (label, value) in enumerate(info_data):
            row = info_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
        
        # Disclaimer
        self.doc.add_paragraph()
        disclaimer = self.doc.add_paragraph()
        disclaimer.add_run("Note: ").font.bold = True
        disclaimer.add_run(
            "Ce document a été généré automatiquement par le système Digital Humans. "
            "Une revue humaine est recommandée avant utilisation en production."
        )
        disclaimer.runs[1].font.italic = True
        disclaimer.runs[1].font.color.rgb = GRAY_TEXT
    
    # ========================================================================
    # HELPERS
    # ========================================================================
    
    def _style_table(self, table, has_header: bool = True):
        """Applique le style aux tableaux"""
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                # Style des cellules
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(3)
                    paragraph.paragraph_format.space_after = Pt(3)
                    for run in paragraph.runs:
                        run.font.size = SIZE_SMALL
                
                # Header row
                if has_header and i == 0:
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{TABLE_HEADER_BG}"/>')
                    cell._tc.get_or_add_tcPr().append(shading)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)
                # Alternating rows
                elif i % 2 == 0:
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{TABLE_ALT_ROW}"/>')
                    cell._tc.get_or_add_tcPr().append(shading)
    
    def _parse_markdown_to_docx(self, markdown_content: str):
        """Convertit du Markdown en éléments Word"""
        lines = markdown_content.split("\n")
        in_code_block = False
        code_buffer = []
        in_table = False
        table_rows = []
        
        for line in lines:
            # Code blocks
            if line.strip().startswith('```'):
                if in_code_block:
                    # Fin du bloc de code
                    self._add_code_block("\n".join(code_buffer))
                    code_buffer = []
                    in_code_block = False
                else:
                    in_code_block = True
                continue
            
            if in_code_block:
                code_buffer.append(line)
                continue
            
            # Tables
            if "|" in line and line.strip().startswith("|"):
                if "---" in line:
                    continue  # Skip separator
                table_rows.append([c.strip() for c in line.split("|") if c.strip()])
                in_table = True
                continue
            elif in_table and table_rows:
                # Fin de table
                self._add_markdown_table(table_rows)
                table_rows = []
                in_table = False
            
            # Headers
            if line.startswith("### "):
                self.doc.add_heading(line[4:].strip(), 3)
            elif line.startswith("## "):
                self.doc.add_heading(line[3:].strip(), 2)
            elif line.startswith("# "):
                self.doc.add_heading(line[2:].strip(), 1)
            # Lists
            elif line.strip().startswith("- **") or line.strip().startswith("* **"):
                # Liste avec texte en gras
                p = self.doc.add_paragraph(style='List Bullet')
                # Extraire le texte gras et le reste
                match = re.match(r'[-*]\s+\*\*(.+?)\*\*:?\s*(.*)', line.strip())
                if match:
                    run = p.add_run(match.group(1))
                    run.font.bold = True
                    if match.group(2):
                        p.add_run(": " + match.group(2))
                else:
                    p.add_run(line.strip()[2:])
            elif line.strip().startswith("- ") or line.strip().startswith("* "):
                p = self.doc.add_paragraph(style='List Bullet')
                text = line.strip()[2:]
                self._add_formatted_text(p, text)
            # Bold text paragraphs
            elif line.strip():
                p = self.doc.add_paragraph()
                self._add_formatted_text(p, line)
        
        # Flush remaining table
        if table_rows:
            self._add_markdown_table(table_rows)
    
    def _add_formatted_text(self, paragraph, text: str):
        """Ajoute du texte avec formatage (gras, italique)"""
        # Pattern pour **bold** et *italic*
        parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
        
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.font.bold = True
            elif part.startswith("*") and part.endswith("*"):
                run = paragraph.add_run(part[1:-1])
                run.font.italic = True
            else:
                paragraph.add_run(part)
    
    def _add_code_block(self, code: str):
        """Ajoute un bloc de code"""
        for line in code.split("\n"):
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            run = p.add_run(line)
            run.font.name = FONT_CODE
            run.font.size = SIZE_CODE
    
    def _add_markdown_table(self, rows: List[List[str]]):
        """Ajoute un tableau depuis des données Markdown"""
        if not rows:
            return
        
        num_cols = max(len(row) for row in rows)
        table = self.doc.add_table(rows=len(rows), cols=num_cols)
        table.style = 'Table Grid'
        self._style_table(table, has_header=True)
        
        for i, row_data in enumerate(rows):
            for j, cell_data in enumerate(row_data):
                if j < num_cols:
                    table.rows[i].cells[j].text = cell_data
    
    def _mermaid_to_image(self, mermaid_code: str) -> Optional[bytes]:
        """Convertit un diagramme Mermaid en image PNG via Kroki"""
        try:
            import requests
            
            # Utiliser Kroki.io pour la conversion
            response = requests.post(
                "https://kroki.io/mermaid/png",
                data=mermaid_code.encode('utf-8'),
                headers={"Content-Type": "text/plain"},
                timeout=30
            )
            
            if response.status_code == 200:
                return response.content
            else:
                logger.warning(f"Kroki returned {response.status_code}: {response.text[:200]}")
                return None
                
        except Exception as e:
            logger.warning(f"Failed to convert Mermaid to image: {e}")
            return None


# ============================================================================
# FONCTION HELPER POUR UTILISATION SIMPLE
# ============================================================================

def generate_sds_docx_v3(
    project_name: str,
    synthesis_result: Dict[str, Any],
    wbs_data: Optional[Dict[str, Any]] = None,
    project_info: Optional[Dict[str, Any]] = None,
    output_path: Optional[str] = None
) -> bytes:
    """
    Génère un document SDS DOCX à partir du résultat de synthèse
    
    Usage:
        from app.services.sds_docx_generator_v3 import generate_sds_docx_v3
        
        docx_bytes = generate_sds_docx_v3(
            project_name="Mon Projet",
            synthesis_result=synthesis_service_result,
            output_path="/tmp/sds.docx"
        )
    """
    generator = SDSDocxGeneratorV3()
    return generator.generate(
        project_name=project_name,
        synthesis_result=synthesis_result,
        wbs_data=wbs_data,
        project_info=project_info,
        output_path=output_path
    )


# Singleton
_generator_instance = None

def get_sds_docx_generator() -> SDSDocxGeneratorV3:
    """Retourne une instance du générateur"""
    global _generator_instance
    if _generator_instance is None:
        _generator_instance = SDSDocxGeneratorV3()
    return _generator_instance
